import pythoncom
import win32com.client
import time
from datetime import datetime
import _strptime
from pythoncom import com_error
import robot.libraries.Screenshot as screenshot
import os
from robot.api import logger
import sys
import ast
import re
import pandas as pd
import openpyxl 
from openpyxl import Workbook
from openpyxl import load_workbook
import json
import docx
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm, Pt, Mm, Inches
from docx.enum.section import WD_ORIENT
import  logging
logging.basicConfig(level=logging.INFO)
import glob 
from openpyxl import load_workbook, Workbook
import win32com.client as win32

import pandas as pd
import docx
# from spire.doc import*
#from spire.doc.common import*
from docx.enum.section import WD_ORIENT
from docx2pdf import convert
from PIL import Image
import shutil
import os
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, RGBColor, Pt
from docx.enum.section import WD_ORIENT
from docx2pdf import convert
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import datetime
from openpyxl.styles import Font, PatternFill
from openpyxl.utils.dataframe import dataframe_to_rows


class SAP_Tcode_Library:
    """The SapGuiLibrary is a library that enables users to create tests for the Sap Gui application

    The library uses the Sap Scripting Engine, therefore Scripting must be enabled in Sap in order for this library to work.

    = Opening a connection / Before running tests =

    First of all, you have to *make sure the Sap Logon Pad is started*. You can automate this process by using the
    AutoIT library or the Process Library.

    After the Sap Login Pad is started, you can connect to the Sap Session using the keyword `connect to session`.

    If you have a successful connection you can use `Open Connection` to open a new connection from the Sap Logon Pad
    or `Connect To Existing Connection` to connect to a connection that is already open.

    = Locating or specifying elements =

    You need to specify elements starting from the window ID, for example, wnd[0]/tbar[1]/btn[8]. In some cases the SAP
    ID contains backslashes. Make sure you escape these backslashes by adding another backslash in front of it.

    = Screenshots (on error) =

    The SapGUILibrary offers an option for automatic screenshots on error.
    Default this option is enabled, use keyword `disable screenshots on error` to skip the screenshot functionality.
    Alternatively, this option can be set at import.
    """
    __version__ = '1.1'
    ROBOT_LIBRARY_SCOPE = 'GLOBAL'

    def __init__(self, screenshots_on_error=True, screenshot_directory=None):
        """Sets default variables for the library
        """
        self.explicit_wait = float(0.0)

        self.sapapp = -1
        self.session = -1
        self.connection = -1

        self.take_screenshots = screenshots_on_error
        self.screenshot = screenshot.Screenshot()

        if screenshot_directory is not None:
            if not os.path.exists(screenshot_directory):
                os.makedirs(screenshot_directory)
            self.screenshot.set_screenshot_directory(screenshot_directory)

    def click_element(self, element_id):
        """Performs a single click on a given element. Used only for buttons, tabs and menu items.

        In case you want to change a value of an element like checkboxes of selecting an option in dropdown lists,
        use `select checkbox` or `select from list by label` instead.
        """

        # Performing the correct method on an element, depending on the type of element
        element_type = self.get_element_type(element_id)
        if (element_type == "GuiTab"
                or element_type == "GuiMenu"):
            self.session.findById(element_id).select()
        elif element_type == "GuiButton":
            self.session.findById(element_id).press()
        else:
            self.take_screenshot()
            message = "You cannot use 'click_element' on element type '%s', maybe use 'select checkbox' instead?" % element_type
            raise Warning(message)
        time.sleep(self.explicit_wait)

    def click_toolbar_button(self, table_id, button_id):
        """Clicks a button of a toolbar within a GridView 'table_id' which is contained within a shell object.
        Use the Scripting tracker recorder to find the 'button_id' of the button to click
        """
        self.element_should_be_present(table_id)

        try:
            self.session.findById(table_id).pressToolbarButton(button_id)
        except AttributeError:
            self.take_screenshot()
            self.session.findById(table_id).pressButton(button_id)
        except com_error:
            self.take_screenshot()
            message = "Cannot find Button_id '%s'." % button_id
            raise ValueError(message)
        time.sleep(self.explicit_wait)

    def connect_to_existing_connection(self, connection_name):
        """Connects to an open connection. If the connection matches the given connection_name, the session is connected
        to this connection.
        """
        self.connection = self.sapapp.Children(0)
        if self.connection.Description == connection_name:
            self.session = self.connection.children(0)
        else:
            self.take_screenshot()
            message = "No existing connection for '%s' found." % connection_name
            raise ValueError(message)

    def connect_to_session(self, explicit_wait=0):
        """Connects to an open session SAP.

        See `Opening a connection / Before running tests` for details about requirements before connecting to a session.

        Optionally `set explicit wait` can be used to set the explicit wait time.

        *Examples*:
        | *Keyword*             | *Attributes*          |
        | connect to session    |                       |
        | connect to session    | 3                     |
        | connect to session    | explicit_wait=500ms   |

        """
        lenstr = len("SAPGUI")
        rot = pythoncom.GetRunningObjectTable()
        rotenum = rot.EnumRunning()
        while True:
            monikers = rotenum.Next()
            if not monikers:
                break
            ctx = pythoncom.CreateBindCtx(0)
            name = monikers[0].GetDisplayName(ctx, None);

            if name[-lenstr:] == "SAPGUI":
                obj = rot.GetObject(monikers[0])
                sapgui = win32com.client.Dispatch(obj.QueryInterface(pythoncom.IID_IDispatch))
                self.sapapp = sapgui.GetScriptingEngine
                # Set explicit_wait after connection succeed
                self.set_explicit_wait(explicit_wait)

        if hasattr(self.sapapp, "OpenConnection") == False:
            self.take_screenshot()
            message = "Could not connect to Session, is Sap Logon Pad open?"
            raise Warning(message)
        # run explicit wait last
        time.sleep(self.explicit_wait)

    def disable_screenshots_on_error(self):
        """Disables automatic screenshots on error.
        """
        self.take_screenshots = False

    def doubleclick_element(self, element_id, item_id, column_id):
        """Performs a double-click on a given element. Used only for shell objects.
        """

        # Performing the correct method on an element, depending on the type of element
        element_type = self.get_element_type(element_id)
        if element_type == "GuiShell":
            self.session.findById(element_id).doubleClickItem(item_id, column_id)
        else:
            self.take_screenshot()
            message = "You cannot use 'doubleclick element' on element type '%s', maybe use 'click element' instead?" % element_type
            raise Warning(message)
        time.sleep(self.explicit_wait)

    def element_should_be_present(self, element_id, message=None):
        """Checks whether an element is present on the screen.
        """
        try:
            self.session.findById(element_id)
        except com_error:
            self.take_screenshot()
            if message is None:
                message = "Cannot find Element '%s'." % element_id
            raise ValueError(message)

    def element_value_should_be(self, element_id, expected_value, message=None):
        """Checks whether the element value is the same as the expected value.
        The possible expected values depend on the type of element (see usage).

         Usage:
         | *Element type*   | *possible values*                 |
         | textfield        | text                              |
         | label            | text                              |
         | checkbox         | checked / unchecked               |
         | radiobutton      | checked / unchecked               |
         | combobox         | text of the option to be expected |
         """
        element_type = self.get_element_type(element_id)
        actual_value = self.get_value(element_id)

        # Breaking up the different element types so we can check the value the correct way
        if (element_type == "GuiTextField"
                or element_type == "GuiCTextField"
                or element_type == "GuiComboBox"
                or element_type == "GuiLabel"):
            self.session.findById(element_id).setfocus()
            time.sleep(self.explicit_wait)
            # In these cases we can simply check the text value against the value of the element
            if expected_value != actual_value:
                if message is None:
                    message = "Element value of '%s' should be '%s', but was '%s'" % (
                        element_id, expected_value, actual_value)
                self.take_screenshot()
                raise AssertionError(message)
        elif element_type == "GuiStatusPane":
            if expected_value != actual_value:
                if message is None:
                    message = "Element value of '%s' should be '%s', but was '%s'" % (
                        element_id, expected_value, actual_value)
                self.take_screenshot()
                raise AssertionError(message)
        elif (element_type == "GuiCheckBox"
              or element_type == "GuiRadioButton"):
            # First check if there is a correct value given, otherwise raise an assertion error
            self.session.findById(element_id).setfocus()
            if (expected_value.lower() != "checked"
                    and expected_value.lower() != "unchecked"):
                # Raise an AsertionError when no correct expected_value is given
                self.take_screenshot()
                if message is None:
                    message = "Incorrect value for element type '%s', provide checked or unchecked" % element_type
                raise AssertionError(message)

            # Check whether the expected value matches the actual value. If not, raise an assertion error
            if expected_value.lower() != actual_value:
                self.take_screenshot()
                if message is None:
                    message = "Element value of '%s' didn't match the expected value" % element_id
                raise AssertionError(message)
        else:
            # When the type of element can't be checked, raise an assertion error
            self.take_screenshot()
            message = "Cannot use keyword 'element value should be' for element type '%s'" % element_type
            raise Warning(message)
        # Run explicit wait as last
        time.sleep(self.explicit_wait)

    def element_value_should_contain(self, element_id, expected_value, message=None):
        """Checks whether the element value contains the expected value.
        The possible expected values depend on the type of element (see usage).

         Usage:
         | *Element type*   | *possible values*                 |
         | textfield        | text                              |
         | label            | text                              |
         | combobox         | text of the option to be expected |
         """
        element_type = self.get_element_type(element_id)

        # Breaking up the different element types so we can check the value the correct way
        if (element_type == "GuiTextField"
                or element_type == "GuiCTextField"
                or element_type == "GuiComboBox"
                or element_type == "GuiLabel"):
            self.session.findById(element_id).setfocus()
            actual_value = self.get_value(element_id)
            time.sleep(self.explicit_wait)
            # In these cases we can simply check the text value against the value of the element
            if expected_value not in actual_value:
                self.take_screenshot()
                if message is None:
                    message = "Element value '%s' does not contain '%s', (but was '%s')" % (
                        element_id, expected_value, actual_value)
                raise AssertionError(message)
        else:
            # When the element content can't be checked, raise an assertion error
            self.take_screenshot()
            message = "Cannot use keyword 'element value should contain' for element type '%s'" % element_type
            raise Warning(message)
        # Run explicit wait as last
        time.sleep(self.explicit_wait)

    def enable_screenshots_on_error(self):
        """Enables automatic screenshots on error.
        """
        self.take_screenshots = True

    def get_table_cell_text(self, table_id, row, column):
        """Returns the cell value for the specified cell.
        """               
  
        try:
            #Access the table control
            table_control = self.session.findById(table_id)
            # Use getCell to access the specific cell
            cell = table_control.getCell(row, column)
            # Get the text from the cell
            cell_text = cell.Text  # Or use other appropriate property or method
            return cell_text
        except com_error:
            self.take_screenshot()
            message = "Cannot find Column_id '%s'." % column
            raise ValueError(message)
        
    def find_all_rows_by_cell_content(self, table_id, column_index, content): 
        print(table_id)
        found_rows = []
        table_control = self.session.findById(table_id)
        try:
            for row_index in range(table_control.RowCount):
                print(row_index)
                cell = table_control.getCell(row_index, column_index)
                print(cell.Text)
                if content in cell.Text:
                    found_rows.append(row_index)
                    if len(found_rows) == 1:
                        found_row = found_rows[0]
                        return found_row
                    elif len(found_rows) > 1:
                        raise ValueError("Array contains more than one value")
                    else:
                        raise ValueError("Array is empty")
        except ValueError as e:
            print("Error while searching table: {e}")

    
    
    def get_cell_value(self, table_id, row_num, col_id):
        """Returns the cell value for the specified cell.
        """
        self.element_should_be_present(table_id)

        try:
            cellValue = self.session.findById(table_id).getCellValue(row_num, col_id)
            return cellValue
        except com_error:
            self.take_screenshot()
            message = "Cannot find Column_id '%s'." % col_id
            raise ValueError(message)

    def get_element_location(self, element_id):
        """Returns the Sap element location for the given element.
        """
        self.element_should_be_present(element_id)
        screenleft = self.session.findById(element_id).screenLeft
        screentop = self.session.findById(element_id).screenTop
        return screenleft, screentop

    def get_element_type(self, element_id):
        """Returns the Sap element type for the given element.
        """
        try:
            type = self.session.findById(element_id).type
            return type
        except com_error:
            self.take_screenshot()
            message = "Cannot find element with id '%s'" % element_id
            raise ValueError(message)

    def get_row_count(self, table_id):
        """Returns the number of rows found in the specified table.
        """
        self.element_should_be_present(table_id)
        rowCount = self.session.findById(table_id).rowCount
        return rowCount

    def get_scroll_position(self, element_id):
        """Returns the scroll position of the scrollbar of an element 'element_id' that is contained within a shell object.
        """
        self.element_should_be_present(element_id)
        currentPosition = self.session.findById(element_id).verticalScrollbar.position
        return currentPosition

    def get_value(self, element_id):
        """Gets the value of the given element. The possible return values depend on the type of element (see Return values).

        Return values:
        | *Element type*   | *Return values*                   |
        | textfield        | text                              |
        | label            | text                              |
        | checkbox         | checked / unchecked               |
        | radiobutton      | checked / unchecked               |
        | combobox         | text of the selected option       |
        | guibutton        | text                              |
        | guititlebar      | text                              |
        | guistatusbar     | text                              |
        | guitab           | text                              |
        """
        element_type = self.get_element_type(element_id)
        return_value = ""
        if (element_type == "GuiTextField"
                or element_type == "GuiCTextField"
                or element_type == "GuiLabel"
                or element_type == "GuiTitlebar"
                or element_type == "GuiStatusbar"
                or element_type == "GuiButton"
                or element_type == "GuiTab"
                or element_type == "GuiShell"):
            self.set_focus(element_id)
            return_value = self.session.findById(element_id).text
        elif element_type == "GuiStatusPane":
            return_value = self.session.findById(element_id).text
        elif (element_type == "GuiCheckBox"
              or element_type == "GuiRadioButton"):
            actual_value = self.session.findById(element_id).selected
            # In these situations we return check / unchecked, so we change these values here
            if actual_value == True:
                return_value = "checked"
            elif actual_value == False:
                return_value = "unchecked"
        elif element_type == "GuiComboBox":
            return_value = self.session.findById(element_id).text
            # In comboboxes there are many spaces after the value. In order to check the value, we strip them away.
            return_value = return_value.strip()
        else:
            # If we can't return the value for this element type, raise an assertion error
            self.take_screenshot()
            message = "Cannot get value for element type '%s'" % element_type
            raise Warning(message)
        return return_value

    def get_window_title(self, locator):
        """Retrieves the window title of the given window.
        """
        return_value = ""
        try:
            return_value = self.session.findById(locator).text
        except com_error:
            self.take_screenshot()
            message = "Cannot find window with locator '%s'" % locator
            raise ValueError(message)

        return return_value

    def input_password(self, element_id, password):
        """Inserts the given password into the text field identified by locator.
        The password is not recorded in the log.
        """
        element_type = self.get_element_type(element_id)
        if (element_type == "GuiTextField"
                or element_type == "GuiCTextField"
                or element_type == "GuiShell"
                or element_type == "GuiPasswordField"):
            self.session.findById(element_id).text = password
            logger.info("Typing password into text field '%s'." % element_id)
            time.sleep(self.explicit_wait)
        else:
            self.take_screenshot()
            message = "Cannot use keyword 'input password' for element type '%s'" % element_type
            raise ValueError(message)

    def input_text(self, element_id, text):
        """Inserts the given text into the text field identified by locator.
        Use keyword `input password` to insert a password in a text field.
        """
        element_type = self.get_element_type(element_id)
        if (element_type == "GuiTextField"
                or element_type == "GuiCTextField"
                or element_type == "GuiShell"
                or element_type == "GuiPasswordField"):
            self.session.findById(element_id).text = text
            logger.info("Typing text '%s' into text field '%s'." % (text, element_id))
            time.sleep(self.explicit_wait)
        else:
            self.take_screenshot()
            message = "Cannot use keyword 'input text' for element type '%s'" % element_type
            raise ValueError(message)

    def maximize_window(self, window=0):
        """Maximizes the SapGui window.
        """
        try:
            self.session.findById("wnd[%s]" % window).maximize()
            time.sleep(self.explicit_wait)
        except com_error:
            self.take_screenshot()
            message = "Cannot maximize window wnd[% s], is the window actually open?" % window
            raise ValueError(message)

        # run explicit wait last
        time.sleep(self.explicit_wait)

    def open_connection(self, connection_name):
        """Opens a connection to the given connection name. Be sure to provide the full connection name, including the bracket part.
        """
        # First check if the sapapp is set and OpenConnection method exists
        if hasattr(self.sapapp, "OpenConnection") == False:
            self.take_screenshot()
            message = "Cannot find an open Sap Login Pad, is Sap Logon Pad open?"
            raise Warning(message)

        try:
            self.connection = self.sapapp.OpenConnection(connection_name, True)
        except com_error:
            self.take_screenshot()
            message = "Cannot open connection '%s', please check connection name." % connection_name
            raise ValueError(message)
        self.session = self.connection.children(0)
        # run explicit wait last
        time.sleep(self.explicit_wait)

    def run_transaction(self, transaction):
        """Runs a Sap transaction. An error is given when an unknown transaction is specified.
        """
        self.session.findById("wnd[0]/tbar[0]/okcd").text = transaction
        time.sleep(self.explicit_wait)
        self.send_vkey(0)

        if transaction == '/nex':
            return

        pane_value = self.session.findById("wnd[0]/sbar/pane[0]").text
        if pane_value in ("Transactie %s bestaat niet" % transaction.upper(),
                          "Transaction %s does not exist" % transaction.upper(),
                          "Transaktion %s existiert nicht" % transaction.upper()):
            self.take_screenshot()
            message = "Unknown transaction: '%s'" % transaction
            raise ValueError(message)

    def scroll(self, element_id, position):
        """Scrolls the scrollbar of an element 'element_id' that is contained within a shell object.
        'Position' is the number of rows to scroll.
        """
        self.element_should_be_present(element_id)
        self.session.findById(element_id).verticalScrollbar.position = position
        time.sleep(self.explicit_wait)

    def select_checkbox(self, element_id):
        """Selects checkbox identified by locator.
        Does nothing if the checkbox is already selected.
        """
        element_type = self.get_element_type(element_id)
        if element_type == "GuiCheckBox":
            self.session.findById(element_id).selected = True
        else:
            self.take_screenshot()
            message = "Cannot use keyword 'select checkbox' for element type '%s'" % element_type
            raise ValueError(message)
        time.sleep(self.explicit_wait)

    def select_context_menu_item(self, element_id, menu_or_button_id, item_id):
        """Selects an item from the context menu by clicking a button or right-clicking in the node context menu.
        """
        self.element_should_be_present(element_id)

        # The function checks if the element has an attribute "nodeContextMenu" or "pressContextButton"
        if hasattr(self.session.findById(element_id), "nodeContextMenu"):
            self.session.findById(element_id).nodeContextMenu(menu_or_button_id)
        elif hasattr(self.session.findById(element_id), "pressContextButton"):
            self.session.findById(element_id).pressContextButton(menu_or_button_id)
        # The element has neither attributes, give an error message
        else:
            self.take_screenshot()
            element_type = self.get_element_type(element_id)
            message = "Cannot use keyword 'select context menu item' for element type '%s'" % element_type
            raise ValueError(message)
        self.session.findById(element_id).selectContextMenuItem(item_id)
        time.sleep(self.explicit_wait)

    def select_from_list_by_label(self, element_id, value):
        """Selects the specified option from the selection list.
        """
        element_type = self.get_element_type(element_id)
        if element_type == "GuiComboBox":
            self.session.findById(element_id).value = value
            time.sleep(self.explicit_wait)
        else:
            self.take_screenshot()
            message = "Cannot use keyword 'select from list by label' for element type '%s'" % element_type
            raise ValueError(message)

    def select_node(self, tree_id, node_id, expand=False):
        """Selects a node of a TableTreeControl 'tree_id' which is contained within a shell object.

        Use the Scripting tracker recorder to find the 'node_id' of the node.
        Expand can be set to True to expand the node. If the node cannot be expanded, no error is given.
        """
        self.element_should_be_present(tree_id)
        self.session.findById(tree_id).selectedNode = node_id
        if expand:
            #TODO: elegantere manier vinden om dit af te vangen
            try:
                self.session.findById(tree_id).expandNode(node_id)
            except com_error:
                pass
        time.sleep(self.explicit_wait)

    def select_node_link(self, tree_id, link_id1, link_id2):
        """Selects a link of a TableTreeControl 'tree_id' which is contained within a shell object.

        Use the Scripting tracker recorder to find the 'link_id1' and 'link_id2' of the link to select.
        """
        self.element_should_be_present(tree_id)
        self.session.findById(tree_id).selectItem(link_id1, link_id2)
        self.session.findById(tree_id).clickLink(link_id1, link_id2)
        time.sleep(self.explicit_wait)

    def select_radio_button(self, element_id):
        """Sets radio button to the specified value.
        """
        element_type = self.get_element_type(element_id)
        if element_type == "GuiRadioButton":
            self.session.findById(element_id).selected = True
        else:
            self.take_screenshot()
            message = "Cannot use keyword 'select radio button' for element type '%s'" % element_type
            raise ValueError(message)
        time.sleep(self.explicit_wait)

    def select_table_column(self, table_id, column_id):
        """Selects an entire column of a GridView 'table_id' which is contained within a shell object.

        Use the Scripting tracker recorder to find the 'column_id' of the column to select.
        """
        self.element_should_be_present(table_id)
        try:
            self.session.findById(table_id).selectColumn(column_id)
        except com_error:
            self.take_screenshot()
            message = "Cannot find Column_id '%s'." % column_id
            raise ValueError(message)
        time.sleep(self.explicit_wait)

    def select_table_row(self, table_id, row_num):
        """Selects an entire row of a table. This can either be a TableControl or a GridView 'table_id'
        which is contained within a shell object. The row is an index to select the row, starting from 0.
        """
        element_type = self.get_element_type(table_id)
        if (element_type == "GuiTableControl"):
            id = self.session.findById(table_id).getAbsoluteRow(row_num)
            id.selected = -1
        else:
            try:
                self.session.findById(table_id).selectedRows = row_num
            except com_error:
                self.take_screenshot()
                message = "Cannot use keyword 'select table row' for element type '%s'" % element_type
                raise ValueError(message)
        time.sleep(self.explicit_wait)

    def send_vkey(self, vkey_id, window=0):
        """Sends a SAP virtual key combination to the window, not into an element.
        If you want to send a value to a text field, use `input text` instead.

        To send a vkey, you can either use te *VKey ID* or the *Key combination*.

        Sap Virtual Keys (on Windows)
        | *VKey ID* | *Key combination*     | *VKey ID* | *Key combination*     | *VKey ID* | *Key combination*     |
        | *0*       | Enter                 | *26*      | Ctrl + F2             | *72*      | Ctrl + A              |
        | *1*       | F1                    | *27*      | Ctrl + F3             | *73*      | Ctrl + D              |
        | *2*       | F2                    | *28*      | Ctrl + F4             | *74*      | Ctrl + N              |
        | *3*       | F3                    | *29*      | Ctrl + F5             | *75*      | Ctrl + O              |
        | *4*       | F4                    | *30*      | Ctrl + F6             | *76*      | Shift + Del           |
        | *5*       | F5                    | *31*      | Ctrl + F7             | *77*      | Ctrl + Ins            |
        | *6*       | F6                    | *32*      | Ctrl + F8             | *78*      | Shift + Ins           |
        | *7*       | F7                    | *33*      | Ctrl + F9             | *79*      | Alt + Backspace       |
        | *8*       | F8                    | *34*      | Ctrl + F10            | *80*      | Ctrl + Page Up        |
        | *9*       | F9                    | *35*      | Ctrl + F11            | *81*      | Page Up               |
        | *10*      | F10                   | *36*      | Ctrl + F12            | *82*      | Page Down             |
        | *11*      | F11 or Ctrl + S       | *37*      | Ctrl + Shift + F1     | *83*      | Ctrl + Page Down      |
        | *12*      | F12 or ESC            | *38*      | Ctrl + Shift + F2     | *84*      | Ctrl + G              |
        | *14*      | Shift + F2            | *39*      | Ctrl + Shift + F3     | *85*      | Ctrl + R              |
        | *15*      | Shift + F3            | *40*      | Ctrl + Shift + F4     | *86*      | Ctrl + P              |
        | *16*      | Shift + F4            | *41*      | Ctrl + Shift + F5     | *87*      | Ctrl + B              |
        | *17*      | Shift + F5            | *42*      | Ctrl + Shift + F6     | *88*      | Ctrl + K              |
        | *18*      | Shift + F6            | *43*      | Ctrl + Shift + F7     | *89*      | Ctrl + T              |
        | *19*      | Shift + F7            | *44*      | Ctrl + Shift + F8     | *90*      | Ctrl + Y              |
        | *20*      | Shift + F8            | *45*      | Ctrl + Shift + F9     | *91*      | Ctrl + X              |
        | *21*      | Shift + F9            | *46*      | Ctrl + Shift + F10    | *92*      | Ctrl + C              |
        | *22*      | Ctrl + Shift + 0      | *47*      | Ctrl + Shift + F11    | *93*      | Ctrl + V              |
        | *23*      | Shift + F11           | *48*      | Ctrl + Shift + F12    | *94*      | Shift + F10           |
        | *24*      | Shift + F12           | *70*      | Ctrl + E              | *97*      | Ctrl + #              |
        | *25*      | Ctrl + F1             | *71*      | Ctrl + F              |           |                       |

        Examples:
        | *Keyword*     | *Attributes*      |           |
        | send_vkey     | 8                 |           |
        | send_vkey     | Ctrl + Shift + F1 |           |
        | send_vkey     | Ctrl + F7         | window=1  |
        """
        vkey_id = str(vkey_id)
        vkeys_array = ["ENTER", "F1", "F2", "F3", "F4", "F5", "F6", "F7", "F8", "F9", "F10", "F11", "F12",
                       None, "SHIFT+F2", "SHIFT+F3", "SHIFT+F4", "SHIFT+F5", "SHIFT+F6", "SHIFT+F7", "SHIFT+F8",
                       "SHIFT+F9", "CTRL+SHIFT+0", "SHIFT+F11", "SHIFT+F12", "CTRL+F1", "CTRL+F2", "CTRL+F3", "CTRL+F4",
                       "CTRL+F5", "CTRL+F6", "CTRL+F7", "CTRL+F8", "CTRL+F9", "CTRL+F10", "CTRL+F11", "CTRL+F12",
                       "CTRL+SHIFT+F1", "CTRL+SHIFT+F2", "CTRL+SHIFT+F3", "CTRL+SHIFT+F4", "CTRL+SHIFT+F5",
                       "CTRL+SHIFT+F6", "CTRL+SHIFT+F7", "CTRL+SHIFT+F8", "CTRL+SHIFT+F9", "CTRL+SHIFT+F10",
                       "CTRL+SHIFT+F11", "CTRL+SHIFT+F12", None, None, None, None, None, None, None, None, None, None,
                       None, None, None, None, None, None, None, None, None, None, None, "CTRL+E", "CTRL+F", "CTRL+A",
                       "CTRL+D", "CTRL+N", "CTRL+O", "SHIFT+DEL", "CTRL+INS", "SHIFT+INS", "ALT+BACKSPACE",
                       "CTRL+PAGEUP", "PAGEUP",
                       "PAGEDOWN", "CTRL+PAGEDOWN", "CTRL+G", "CTRL+R", "CTRL+P", "CTRL+B", "CTRL+K", "CTRL+T",
                       "CTRL+Y",
                       "CTRL+X", "CTRL+C", "CTRL+V", "SHIFT+F10", None, None, "CTRL+#"]

        # If a key combi is given, replace vkey_id by correct id based on given combination
        if not vkey_id.isdigit():
            search_comb = vkey_id.upper()
            search_comb = search_comb.replace(" ", "")
            search_comb = search_comb.replace("CONTROL", "CTRL")
            search_comb = search_comb.replace("DELETE", "DEL")
            search_comb = search_comb.replace("INSERT", "INS")
            try:
                vkey_id = vkeys_array.index(search_comb)
            except ValueError:
                if search_comb == "CTRL+S":
                    vkey_id = 11
                elif search_comb == "ESC":
                    vkey_id = 12
                else:
                    message = "Cannot find given Vkey, provide a valid Vkey number or combination"
                    raise ValueError(message)

        try:
            self.session.findById("wnd[% s]" % window).sendVKey(vkey_id)
        except com_error:
            self.take_screenshot()
            message = "Cannot send Vkey to given window, is window wnd[% s] actually open?" % window
            raise ValueError(message)
        time.sleep(self.explicit_wait)

    def set_cell_value(self, table_id, row_num, col_id, text):
        """Sets the cell value for the specified cell of a GridView 'table_id' which is contained within a shell object.

        Use the Scripting tracker recorder to find the 'col_id' of the cell to set.
        """
        self.element_should_be_present(table_id)

        try:
            self.session.findById(table_id).modifyCell(row_num, col_id, text)
            logger.info("Typing text '%s' into cell '%s', '%s'" % (text, row_num, col_id))
            time.sleep(self.explicit_wait)
        except com_error:
            self.take_screenshot()
            message = "Cannot type text '%s' into cell '%s', '%s'" % (text, row_num, col_id)
            raise ValueError(message)

    def set_explicit_wait(self, speed):
        """Sets the delay time that is waited after each SapGui keyword.

        The value can be given as a number that is considered to be seconds or as a human-readable string like 1 second
        or 700 ms.

        This functionality is designed to be used for demonstration and debugging purposes. It is not advised to use
        this keyword to wait for an element to appear or function to finish.

         *Possible time formats:*
        | miliseconds       | milliseconds, millisecond, millis, ms |
        | seconds           | seconds, second, secs, sec, s         |
        | minutes           | minutes, minute, mins, min, m         |

         *Example:*
        | *Keyword*         | *Attributes*      |
        | Set explicit wait | 1                 |
        | Set explicit wait | 3 seconds         |
        | Set explicit wait | 500 ms            |
        """
        speed = str(speed)
        if not speed.isdigit():
            speed_elements = speed.split()
            if not speed_elements[0].isdigit():
                message = "The given speed %s doesn't begin with an numeric value, but it should" % speed
                raise ValueError(message)
            else:
                speed_elements[0] = float(speed_elements[0])
                speed_elements[1] = speed_elements[1].lower()
                if (speed_elements[1] == "seconds"
                        or speed_elements[1] == "second"
                        or speed_elements[1] == "s"
                        or speed_elements[1] == "secs"
                        or speed_elements[1] == "sec"):
                    self.explicit_wait = speed_elements[0]
                elif (speed_elements[1] == "minutes"
                      or speed_elements[1] == "minute"
                      or speed_elements[1] == "mins"
                      or speed_elements[1] == "min"
                      or speed_elements[1] == "m"):
                    self.explicit_wait = speed_elements[0] * 60
                elif (speed_elements[1] == "milliseconds"
                      or speed_elements[1] == "millisecond"
                      or speed_elements[1] == "millis"
                      or speed_elements[1] == "ms"):
                    self.explicit_wait = speed_elements[0] / 1000
                else:
                    self.take_screenshot()
                    message = "%s is a unknown time format" % speed_elements[1]
                    raise ValueError(message)
        else:
            # No timeformat given, so time is expected to be given in seconds
            self.explicit_wait = float(speed)

    def set_focus(self, element_id):

        """Sets the focus to the given element.
        """
        element_type = self.get_element_type(element_id)
        if element_type != "GuiStatusPane":
            self.session.findById(element_id).setFocus()
        time.sleep(self.explicit_wait)

    def take_screenshot(self, screenshot_name="sap-screenshot"):
        """Takes a screenshot, only if 'screenshots on error' has been enabled,
        either at import of with keyword `enable screenshots on error`.

        This keyword uses Robots' internal `Screenshot` library.
        """
        if self.take_screenshots == True:
            self.screenshot.take_screenshot(screenshot_name)

    def unselect_checkbox(self, element_id):
        """Removes selection of checkbox identified by locator.
        Does nothing if the checkbox is not selected.
        """
        element_type = self.get_element_type(element_id)
        if element_type == "GuiCheckBox":
            self.session.findById(element_id).selected = False
        else:
            self.take_screenshot()
            message = "Cannot use keyword 'unselect checkbox' for element type '%s'" % element_type
            raise ValueError(message)
        time.sleep(self.explicit_wait)
    
    #New scripts
        
    def is_imp_notes_existing(self, modal_window_id, modal_continue_id):   
        try:
            content = self.session.findById(modal_window_id).Text
            if content == "SAINT: Important SAP Notes":
                print("Modal window exists")
                self.session.findById(modal_continue_id).press()
                return content
            else:
                print("Modal window does not exist.")
            

        except Exception as e:
            print(f"Error: {str(e)}")
            return False
        
    def get_finish_cell_text(self, finish_str, button_id, status_line, refresh_id):
        try:
            while True:
                cell_text_1 = self.session.findById(status_line).Text
                cell_text_2 = cell_text_1[1:]

                if finish_str == cell_text_2:
                    self.session.findById(button_id).press()
                    print("Installation Successful")
                    break  # Exit the loop if the condition is met
                else:
                    self.session.findById(refresh_id).press()
                    #print("No Match")
                    time.sleep(60)

            return cell_text_2
            
        except Exception as e:
            return f"Error: {str(e)}"
            # return cell_text_2
             
    def get_finish_cell_text1(self, finish_str, button_id, status_line, refresh_id):
        try:
            while True:
                cell_text_1 = self.session.findById(status_line).Text
                # cell_text_2 = cell_text_1

                if finish_str == cell_text_1:
                    self.session.findById(button_id).select()
                    print("Installation Successful")
                    break  # Exit the loop if the condition is met
                else:
                    self.session.findById(refresh_id).press()
                    #print("No Match")
                    time.sleep(30)

            return cell_text_1

        except Exception as e:
            return f"Error: {str(e)}"
            # return cell_text_2
 
    def get_maintenance_certificate_text(self, certificate_id):
        try:
            found = False
            while not found:
                license_text = self.session.findById(certificate_id).Text
                license_split = license_text.split()
                license_text_1 = ' '.join(license_split[:-1])
                     
                if license_text_1 == "A valid maintenance certificate was found for system":
                    print("License available to proceed further")
                    found = True    
                else:
                    print("No Valid Maintenance Certificate is found in the System")
                    break
    
        except Exception as e:
            return f"Error: {str(e)}"
        
    def run_time_error_existing(self, runtimeerror_id, back_id):   
        try:
            content = self.session.findById(runtimeerror_id).Text
            if content == "Runtime Error - Description of Exception":
                print("Runtime error exists")
                self.session.findById(back_id).press()
                return content
            else:
                print("Runtime error does not exist.")
            

        except Exception as e:
            print(f"Error: {str(e)}")
            return False
        
    def no_queue_pending(self, no_Queue_id):   
        try:
            content = self.session.findById(no_Queue_id).Text
            if content == "No queue has been defined":
                print("No queue available")
                return content
            else:
                print("Queue is available")
            

        except Exception as e:
            print(f"Error: {str(e)}")
            return False
        
    def import_information(self, text_id, import_id, text1_id, continue_id):  
        try:
            found = False
            while not found:
                importstart_text = self.session.findById(text_id).Text
                importstart_split = importstart_text.split()
                importstart_text_1 = ' '.join(importstart_split[:3] + importstart_split[3:])

                importcomplete_text = self.session.findById(text1_id).Text
                importcomplete_split = importcomplete_text.split()
                importcomplete_text_1 = ' '.join(importcomplete_split[:2] + importcomplete_split[3:])

                if importstart_text_1 == "The SPAM/SAINT update is being imported.":
                    print("Import started")
                    found = True 
                    self.session.findById(import_id).press()
                    return importstart_text_1   
                elif importcomplete_text_1 == "SPAM/SAINT update has already been imported successfully":
                    print("Import completed")
                    found = True
                    self.session.findById(continue_id).press()
                    return importcomplete_text_1

        except Exception as e:
            print(f"Error: {str(e)}")
            return False
        

    def import_success(self, text_id, continue_id):  
        try:
            found = False
            while not found:
                importsuccess_text = self.session.findById(text_id).Text
                importsuccess_split = importsuccess_text.split()
                importsuccess_text_1 = ' '.join(importsuccess_split[:1] + importsuccess_split[2:])

                if importsuccess_text_1 == "Queue SAPK-74002INSTPI imported successfully":
                    print("Import success")
                    found = True 
                    self.session.findById(continue_id).press()
                    return importsuccess_text_1   
                else:
                    print("Import not completed")
                    
        except Exception as e:
            print(f"Error: {str(e)}")
            return False
    
    def epilogue_handling(self, epi_id, spam_id, window_1_id, import_id):   
        try:   
            content = self.session.findById(epi_id).Text
            if content == "EPILOGUE":
                print("patch in Epilogue")
                self.session.findById(spam_id).select()
                content1 = self.session.findById(window_1_id).Text
                if content1 == "phase EPILOGUE.":
                    self.session.findById(import_id).press()
                    return content1
                else:
                    print("phase is not EPILOGUE")
                return content
            
            else:
                print("Queue is not confirmed")

        except Exception as e:
            print(f"Error: {str(e)}")
            return False
        
    def version_print(self, version_id):   
        try:   
            content = self.session.findById(version_id).Text
            content_split = content.split()
            content1 = ' '.join(content_split[:-1])

            if content1 == "Support Package Manager - Version":
                print(content)
                return content
            else:
                print("check spam version manually")

        except Exception as e:
            print(f"Error: {str(e)}")
            return False
         
        
    def confirm(self, content, confirm_id, confirm_queue, refresh1_id):
        try:
            content = self.session.findById(confirm_id).Text

            if content == "Confirm queue":
                self.session.findById(confirm_queue).press()
                print("Confirm queue")
            else:
                self.session.findById(refresh1_id).press()
                #time.sleep(60)

            return content

        except Exception as e:
            print(f"Error: {str(e)}")
            return False

    def is_transaction_locked_by(self, window_id, button_id):
        try:
            lock_text = self.session.findById(window_id).Text
            lock_text_split = lock_text.split()
            lock_text1 = ' '.join(lock_text_split[:-1])
            if lock_text1 == "transaction SPAM is locked by":
                print(f"{lock_text}, so exiting the script")
                self.session.findById(button_id).press()
                self.run_transaction("/nex")
                return lock_text
            else:
                pass
        except Exception as e:
            print(f"Error: {str(e)}")
            return False
    
        
    def spam_search_and_select_label(self, user_area_id, search_text, max_scrolls=50):
        try:
            user_area = self.session.findById(user_area_id)
            scroll_count = 0
            found = False

            while scroll_count < max_scrolls and not found:
                for child in user_area.Children:
                    if child.Text == search_text:
                        print(f"Text Found: {child.Text}")
                        child.SetFocus()
                        self.session.findById("wnd[1]").sendVKey(2)  # Simulate Enter key press
                        found = True
                        break

                if not found:
                    # Scroll down and wait for the content to update
                    print(scroll_count)
                    self.session.findById("wnd[1]").sendVKey(82)  # 86 is the code for Page Down
                    time.sleep(1)  # Adjust as necessary for GUI response time
                    scroll_count += 1

            if not found:
                print("Text not found after scrolling through all pages.")

        except Exception as e:
            print(f"Error: {e}")

    def select_spam_based_on_text(self, control_id, search_text):
        try:
            control = self.session.findById(control_id)
            row_count = control.RowCount  # Assuming the control has a RowCount property
            print(row_count)
            for row in range(row_count):
                print(row)
                cell_value=control
                cell_value = control.GetCellValue(row,"COMPONENT")
                print(cell_value)
                if search_text in cell_value:
                    result = row
                    print("Text Found in ${row}")
                    return row
                else:
                    print("not found")
        except Exception as e:
            return f"Error: {e}"
    
    def spam_multiple_patch_version_select(self, comp_id, search_comp_1, search_patch_1):
        search_comp = ast.literal_eval(search_comp_1)
        search_patch = ast.literal_eval(search_patch_1)
        # print(search_comp, type(search_comp))
        # print(search_patch, type(search_patch))
        # if not len(search_comp) == len(search_patch):
        #     sys.exit() 
        
        comp_area = self.session.FindById(comp_id)
        row_count = comp_area.RowCount

        for i in range(len(search_comp)):
            comp = search_comp[i]
            patch = search_patch[i]
            print(comp, patch)
            
            try:
                for x in range(row_count + 1):
                    cell_value = comp_area.GetCellValue(x, "COMPONENT")
                    print(x, cell_value)
                    if cell_value == comp:
                        comp_area.modifyCell(x,"PATCH_REQ",patch)
                        cell_value_1 = comp_area.GetCellValue(x, "COMPONENT")
                        print("Cell Value 1", cell_value_1)
            except Exception as e:
                 return f"Error: {e}"
    
    def double_click_on_tree_item(self, tree_id, id):
        try:
            tree = self.session.findById(tree_id)
            tree.selectedNode = id
            tree.doubleClickNode(id)
        except Exception as e:
            print("Error: {e}")

    def set_caret_position(self, element_id, position):
        try:
            element = self.session.findById(element_id)
            element.caretposition = position

        except Exception as e:
            print("Error: {e}")
            
    def scot_tree(self, tree_id):
        try:
            tree = self.session.findById(tree_id)
            tree.DoubleClickNode("         23")
    
        except Exception as e:
            print("Error: {e}")

    def select_label(self, user_area_id, search_text, max_scrolls=50):
        try:
            user_area = self.session.findById(user_area_id)
            scroll_count = 0
            found = False

            while scroll_count < max_scrolls and not found:
                for child in user_area.Children:
                    if child.Text == search_text:
                        print(f"Text Found: {child.Text}")
                        child.SetFocus()
                        self.session.findById("wnd[0]").sendVKey(2)  # Simulate Enter key press
                        found = True
                        break

                if not found:
                    # Scroll down and wait for the content to update
                    print(scroll_count)
                    self.session.findById("wnd[0]").sendVKey(82)  # 86 is the code for Page Down
                    # time.sleep(1)  # Adjust as necessary for GUI response time
                    scroll_count += 1

            if not found:
                print("Text not found after scrolling through all pages.")

        except Exception as e:
            print(f"Error: {e}")    

    def selected_rows(self, tree_id, first_visible_row):
        try:
            tree = self.session.findById(tree_id)

            tree.firstVisibleRow = first_visible_row
                

        except Exception as e:
            print(f"Error: {e}")

    def scroll_pagedown(self, window_id):
        try:
            # session.findById("wnd[0]/usr/txtCERT-FPSHA1").setFocus()
            self.session.findById(window_id).setFocus()
            self.session.findById("wnd[0]").TabBackward()
        except Exception as e:
            print(f"Error: {e}")

    def get_grid_ids(self, grid_id):
        try:
            grid_control = self.session.findById(grid_id)

            # Get the number of rows and columns in the grid
            rows = grid_control.RowCount
            columns = grid_control.ColumnCount

            # Retrieve the item IDs and column IDs
            item_ids = [f"{grid_id}/shell[0]/shell[{i}]" for i in range(rows)]
            column_ids = [f"{grid_id}/shell[0]/shell[0]/shell[{i}]" for i in range(columns)]

            return item_ids, column_ids

        except Exception as e:
            print(f"Error: {e}")
            return None, None

    def select_item_from_guilabel(self, user_id, search_text):
        user_area = self.session.findById(user_id)
        labels = [child.Text for child in user_area.Children]
        item_count = user_area.Children.Count
        # print("User Area Labels:", labels, item_count)
        for i in range(item_count):
            element = user_area.Children.ElementAt(i)
            if element.Text.strip() == search_text.strip():
                print(f"Element found: {element.Text}")
                element.SetFocus()
                self.session.findById("wnd[0]").sendVKey(2)
                return
            
    def rows_from_stms(self, table_id): 
        print(table_id)
        row_count = self.session.findById(table_id).rowcount
        print(row_count)
        column_count = self.session.findbyId(table_id).columncount
        print(column_count)
        try:
            for row in range(row_count):
                print(row)
                cell_value_1= self.session.findById(table_id).GetCellValue(row, "SYSNAM")
                self.session.findById(table_id).DoubleClick(row, "SYSNAM")
                print(cell_value_1)
                return cell_value_1
        except Exception as e:
            return f"Error: {e}"  

    def get_cell_value_from_gridtable(self, table_id):
        try:
            control = self.session.findById(table_id)
            row_count = control.RowCount  # Assuming the control has a RowCount property
            col_count = control.ColumnCount
            print(row_count, col_count)
            for row in range(row_count):
                print(row)
                cell_value = control.GetCellValue(row, "DEST")
                print(cell_value)
                return cell_value
        except Exception as e:
            return f"Error: {e}"

    
    def get_no_entries_found_text(self, text_id):
        try:
            found = False
            while not found:
                entry_text = self.session.findById(text_id).Text
                                     
                if entry_text == "No entries found that match selection criteria":
                    print("No entries found that match selection criteria")
                    found = True    
                else:
                    print("Entries are displayed")
                    break
        except Exception as e:
            return f"Error: {e}"
    
    
    def multiple_logon_handling(self, logon_window_id, logon_id, continue_id):  
        try:
            content = self.session.findById(logon_window_id).Text
            if content == "License Information for Multiple Logons":
                print("Multiple logon exists")
                self.session.findById(logon_id).selected = True
                self.session.findById(continue_id).press()
                return content
            else:
                print("Multiple logon does not exist.")
        except Exception as e:
            return f"Error: {e}"      

    def table_scroll(self, table_id, first_visible_row):
        try:
            # tree = self.session.findById(table_id)
            # tree.firstVisibleRow = first_visible_row
            self.session.findById("wnd[0]/usr/cntlGRID1/shellcont/shell").currentCellRow = 29
            self.session.findById("wnd[0]/usr/cntlGRID1/shellcont/shell").firstVisibleRow = 6
            self.session.findById("wnd[0]/usr/cntlGRID1/shellcont/shell").selectedRows = "29"
        except Exception as e:
            print(f"Error: {e}")

    def double_click_on_inside_table(self, tree_id, row_number):
        try:
            tree = self.session.findById(tree_id)
            node = tree.GetNode(row_number)
            node.DoubleClick()
        except Exception as e:
            print(f"Error: {e}")
 
    def double_click_on_current_cell(self, table_id):
        try:
            table = self.session.findById(table_id)
            table.DoubleClickCurrentCell()  # Perform a double-click action on the current cell
 
        except Exception as e:
            print(f"Error: {e}")

    def expand_Element(self, tree_id, node_id):
        try:
            element = self.session.findById(tree_id)
            element.expandNode(f"{node_id}")
        except Exception as e:
            print(f"An error occurred while expanding node: {e}")  

    def select_top_node(self,tree_id, node_id,expand=False):
        self.element_should_be_present(tree_id)
        self.session.findById(tree_id).selectedNode = node_id
        if expand:
            #TODO: elegantere manier vinden om dit af te vangen
            try:
                self.session.findById(tree_id).topNode(node_id)
            except com_error:
                pass
        time.sleep(self.explicit_wait) 

    def expand_node(self, tree_id, node_id):
        element = self.session.findById(tree_id)
        element.expandNode(f"{node_id}")

    def select_item(self, tree_id, nodeid1, nodeid2):
        element=self.session.findById(tree_id)
        element.selectItem(f"{nodeid1}",nodeid2)

    def get_open_items(self, status_id):
        try:
            status = self.session.findById(status_id).Text
            pattern = r"(\d+)\s+items\s+displayed"
            match = re.search(pattern, status)
            if match:
                open_items = match.group(1)
                return open_items
            else:
                print("No match found")
                return None
        except Exception as e:
            return f"Error: {str(e)}"
        
    def set_key_value(self, element_id, key_value):
        self.session.FindById(element_id).key = key_value

    def double_click_table(self, table_id, row, column):
        self.session.FindById(table_id).doubleClickCurrentCell(row, column)

    def search_and_select_lock(self, table_id, lock):
        try:
            table = self.session.FindById(table_id)
            row_count = table.RowCount  # Assuming the control has a RowCount property
            print(row_count)
            for row in range(row_count):
                print(row)
                cell_value=table
                cell_value = table.GetCellValue(row,"GNAME")
                print(cell_value)
                if lock in cell_value:
                    result = row
                    self.session.findById(table_id).selectedRows = row
                    print("Text Found in ${row}")
                    return row
                else:
                    print("not found")
        except Exception as e:
            return f"Error: {e}"
    def window_handling(self, element_id, text, button_id):
        window = self.session.findById(element_id).Text
        if window == text :
            self.session.findById(button_id).press()
            
    def clear_field_text(self, field_id):
        try:
            field = self.session.findById(field_id)
            field.Text = ""
            print(f"Text cleared in field with ID: {field_id}")
        except Exception as e:
            print(f"Error: {e}")
    
    def Excel_Arrange(self, file_location, sheet_name, filename):
        try:
            file_path = f"{file_location}\\{filename}"
            df = pd.read_excel(file_path, sheet_name=sheet_name, header=None)
            df = df.applymap(lambda x: x.strip() if isinstance(x, str) else x)
            df = df.iloc[4:].reset_index(drop=True)
            df.columns = df.iloc[0]
            df = df[1:].reset_index(drop=True)
            df.dropna(how='all', inplace=True)
            df.dropna(axis=1, how='all', inplace=True)
            with pd.ExcelWriter(file_path, engine='openpyxl', mode='a', if_sheet_exists='replace') as writer:
                df.to_excel(writer, sheet_name=sheet_name, index=False)
        except Exception as e:
            pass

    def software_component_version(self, comp_id, search_comp):      
        comp_area = self.session.FindById(comp_id)
        row_count = comp_area.RowCount
        try:
            for x in range(row_count):
                print (x)   
                cell_value = comp_area.GetCellValue(x, "COMPONENT")
                print (cell_value)
                if cell_value == search_comp:
                    version = comp_area.GetCellValue(x, "RELEASE")
                    print(f"Found version for {search_comp}: {version}")
                    return version
                    break  
                else:
                    print(f"Component {search_comp} not found.")
        except Exception as e:
            print(f"Error while searching for {search_comp}: {e}")

    def software_support_package_version(self, comp_id, search_comp):      
        comp_area = self.session.FindById(comp_id)
        row_count = comp_area.RowCount
        try:
            for x in range(row_count):
                # print (x)   
                cell_value = comp_area.GetCellValue(x, "COMPONENT")
                # print (cell_value)
                if cell_value == search_comp:
                    patch = comp_area.GetCellValue(x, "HIGH_PATCH")
                    # print(f"Found version for {search_comp}: {patch}")
                    return patch
                    break  
                else:
                    print(f"Component {search_comp} not found.")
        except Exception as e:
            print(f"Error while searching for {search_comp}: {e}")

    def select_profile_label(self, user_area_id, search_text, max_scrolls=5):
        try:
            user_area = self.session.findById(user_area_id)
            scroll_count = 0
            found = False
 
            while scroll_count < max_scrolls and not found:
                for child in user_area.Children:
                    if child.Text == search_text:
                        print(f"Text Found: {child.Text}")
                        child.SetFocus()
                        # self.session.findById("wnd[1]").sendVKey(2)  # Simulate Enter key press
                        found = True
                        break
 
                if not found:
                    # Scroll down and wait for the content to update
                    print(scroll_count)
                    self.session.findById("wnd[1]").sendVKey(82)  # 86 is the code for Page Down
                    time.sleep(1)  # Adjust as necessary for GUI response time
                    scroll_count += 1
 
            if not found:
                print("Text not found after scrolling through all pages.")
 
        except Exception as e:
            print(f"Error: {e}")

    def check_parameter_found(self, lable_id, parameter):
        user_area = self.session.findById(lable_id)
        item_count = user_area.Children.Count
        for i in range(item_count):
            element = user_area.Children.ElementAt(i)
            if element.Text.strip() == parameter.strip():
                print(element.Text)
                return(element.Text)
        not_found_message = f"Search text {parameter} not found"
        print(f"Search text {parameter} not found")
        return not_found_message
    
    def get_parameter_value(self, lable_id, parameter):
        user_area = self.session.findById(lable_id)
        item_count = user_area.Children.Count
        for i in range(item_count):
            element = user_area.Children.ElementAt(i)
            if element.Text.strip() == parameter.strip():
                element.setFocus()
                self.session.findById("wnd[0]").sendVKey(2)
                return
    
    def manage_window(self, element_id, text, button_id):
        try:
            window_title = self.session.findById(element_id).Text
            window_title_split = window_title.split()
            window = " ".join(window_title_split[:-1])
            if window == text :
                self.session.findById(button_id).press()
            else:
                print(window_title)
        except Exception as e:
            print(f"Error: {e}")

    def double_click_current_cell_value(self, element_id, cell_value):
        try:
            element = self.session.findById(element_id)
            element.currentCellColumn = cell_value
            element.doubleClickCurrentCell()
        except Exception as e:
            print(f"Error: {e}")

    def get_file_content(Self, file_path, ):
        
        with open(file_path, 'r') as file:
            content = file.read()
        return content
    
    def select_org_label(self, user_area_id, search_text, max_scrolls=5):
        try:
            user_area = self.session.findById(user_area_id)
            scroll_count = 0
            found = False

            while scroll_count < max_scrolls and not found:
                for child in user_area.Children:
                    if child.Text == search_text:
                        print(f"Text Found: {child.Text}")
                        child.SetFocus()
                        # self.session.findById("wnd[1]").sendVKey(2)  # Simulate Enter key press
                        found = True
                        break

                if not found:
                    # Scroll down and wait for the content to update
                    print(scroll_count)
                    self.session.findById("wnd[1]").sendVKey(82)  # 86 is the code for Page Down
                    time.sleep(1)  # Adjust as necessary for GUI response time
                    scroll_count += 1

            if not found:
                print("Text not found after scrolling through all pages.")

        except Exception as e:
            print(f"Error: {e}")

    def window_handling(self, window_id, text, button_id):   
        try:   
            content = self.session.findById(window_id).Text
            if content == text:
                print(content)
                self.session.findById(button_id).press()
                return(content)                
            else:
                print(content)
        except Exception as e:
            print(f"Error: {e}")

    def quantity_handling(self, window_id, text, button_id1, button_id2):   
        try:   
            content = self.session.findById(window_id).Text
            if content == text:
                print(content)
                self.session.findById(button_id1).press()
                self.session.findById(button_id2).press()
                return(content)                
            else:
                print(content)
        except Exception as e:
            print(f"Error: {e}")

    def incomplete_log_handle(self, window_id, text1, button_id1, element_id, text2, button_id2):   
        try:   
            content = self.session.findById(window_id).Text
            if content == text1:
                print(content)
                self.session.findById(button_id1).press()
                self.session.findById(element_id).text = text2
                self.session.findById("wnd[0]").sendVKey(0)
                self.session.findById(button_id2).press()
                return(content)                
            else:
                print(content)
        except Exception as e:
            print(f"Error: {e}")
    
    def quantity_select(self, material, quantity, amount, window_id, text, button_id1, button_id2, error_id):
        mat_txt = "wnd[0]/usr/tabsTAXI_TABSTRIP_OVERVIEW/tabpT\\01/ssubSUBSCREEN_BODY:SAPMV45A:4400/subSUBSCREEN_TC:SAPMV45A:4900/tblSAPMV45ATCTRL_U_ERF_AUFTRAG/ctxtRV45A-MABNR"
        qty_txt = "wnd[0]/usr/tabsTAXI_TABSTRIP_OVERVIEW/tabpT\\01/ssubSUBSCREEN_BODY:SAPMV45A:4400/subSUBSCREEN_TC:SAPMV45A:4900/tblSAPMV45ATCTRL_U_ERF_AUFTRAG/txtRV45A-KWMENG"
        amt_txt = "wnd[0]/usr/tabsTAXI_TABSTRIP_OVERVIEW/tabpT\\01/ssubSUBSCREEN_BODY:SAPMV45A:4400/subSUBSCREEN_TC:SAPMV45A:4900/tblSAPMV45ATCTRL_U_ERF_AUFTRAG/txtKOMV-KBETR"
        try:
            for i in range(len(material)):
                mat_id = f"{mat_txt}[1,{i}]"
                print(mat_id)
                qty_id = f"{qty_txt}[3,{i}]"
                print(qty_id)
                amt_id = f"{amt_txt}[15,{i}]"
                print(amt_id)
                print(material[i])
                print(quantity[i])
                print(amount[i])
                self.session.findById(mat_id).text = material[i]
                self.session.findById(qty_id).text = quantity[i]
                self.session.findById(amt_id).text = amount[i]
                self.session.findById("wnd[0]").sendVKey(0)
                self.exceed_quantity_handling(error_id)
                time.sleep(2)
                self.quantity_handling(window_id, text, button_id1, button_id2)
                time.sleep(2)             

        except Exception as e:
            print(e)

    def picked_qty_select(self, picked_qty):
        picked_txt = "wnd[0]/usr/tabsTAXI_TABSTRIP_OVERVIEW/tabpT\\01/ssubSUBSCREEN_BODY:SAPMV50A:1102/tblSAPMV50ATC_LIPS_OVER/txtLIPSD-PIKMG"
        
        try:
            for i in range(len(picked_qty)):
                picked_id = f"{picked_txt}[18,{i}]"
                print(picked_id)
                print(picked_qty[i])
                
                self.session.findById(picked_id).text = picked_qty[i]
                
                self.session.findById("wnd[0]").sendVKey(0)                         

        except Exception as e:
            print(e)



    def exceed_quantity_handling(self, error_id):
        try:
            status = self.session.findById(error_id).text
            print(status)
            status_split = status.split()
            status_splits = status_split[:-1]
            status_text = [status_split for status_split in status_splits if not status_split.isnumeric()]  # Fix the index here
            status1 = ' '.join(status_text)
            print(status1)
            if status1 == "Reorder point for item has been exceeded:":
                found = True
                self.session.findById("wnd[0]").sendVKey(0)
                time.sleep(5)
                return status
            else:
                print(status)  # Fix the variable name here
        except Exception as e:
            print(f"Error: {e}")

    def sales_order_number(self, status_id):
        try:
            status = self.session.findById(status_id).text
            print(status)
            pattern = r'\b\d+\b'
            print(pattern)
            order_numbers = re.findall(pattern, status)
            print(order_numbers)
            if order_numbers:
                order_number = order_numbers[0]
                print("Order Number:", order_number)
                return order_number
            else:
                print("No order number found.")
        except Exception as e:
            print(f"Error: {e}")

    def Outbound_number(self, status_id):
        try:
            # Retrieve the status text
            status = self.session.findById(status_id).text
            print(status)

            # Updated pattern to specifically match the number after "Outbound delivery"
            pattern = r'Outbound delivery (\d+) saved'
            
            # Find all matching numbers
            order_numbers = re.findall(pattern, status)
            print(order_numbers)
            
            if order_numbers:
                # The first match is the order number we're interested in
                order_number = order_numbers[0]
                print("Order Number:", order_number)
                return order_number
            else:
                print("No order number found.")
        except Exception as e:
            print(f"Error: {e}")

    def output_handling(self, window_id, text, label_id, button_id):   
        try:   
            content = self.session.findById(window_id).Text
            if content == text:
                print(content)
                status = self.session.findById(label_id).text
                print(status)
                status_split = status.split()
                status_splits = status_split[:-1]
                status_text = [status_split for status_split in status_splits if not status_split.isnumeric()]  # Fix the index here
                status1 = ' '.join(status_text)
                print(status1)
                if status1 == "IDoc was added and passed for output":
                    found = True
                    self.session.findById(button_id).press()
                    time.sleep(5)
                    return status
                else:
                    print(status)  # Fix the variable name here
            else:
                print(status)
        except Exception as e:
            print(f"Error: {e}")

    def verify_the_idoc_jobs(self, table_id, search_text, process_log_btn, max_attempts=20):
        try:
            control = self.session.findById(table_id)
            row = control.RowCount
            # print(row)
            for i in range(row):
                job_id = f"{table_id}/ctxtDNAST-KSCHL[1,{i}]"
                # print(job_id)
                cell_value = self.session.findById(job_id).Text
                # print(cell_value)
                if cell_value == search_text:
                    status_id = f"{table_id}/lblDV70A-STATUSICON[0,{i}]"
                    status = self.session.findById(status_id).tooltip
                    print(status)
                    if status == "Successfully processed":
                        control.getAbsoluteRow(i).selected = -1
                        self.session.findById(process_log_btn).press()
                        time.sleep(5)
                        self.session.findById("wnd[1]").close()
                        break
                    else:
                        return(status)
                        time.sleep(10)                 
        except Exception as e:
            return f"Error: {e}"
        
    def quantity_select(self, material, quantity, amount):
        mat_txt = "wnd[0]/usr/tabsTAXI_TABSTRIP_OVERVIEW/tabpT\\01/ssubSUBSCREEN_BODY:SAPMV45A:4400/subSUBSCREEN_TC:SAPMV45A:4900/tblSAPMV45ATCTRL_U_ERF_AUFTRAG/ctxtRV45A-MABNR"
        qty_txt = "wnd[0]/usr/tabsTAXI_TABSTRIP_OVERVIEW/tabpT\\01/ssubSUBSCREEN_BODY:SAPMV45A:4400/subSUBSCREEN_TC:SAPMV45A:4900/tblSAPMV45ATCTRL_U_ERF_AUFTRAG/txtRV45A-KWMENG"
        amt_txt = "wnd[0]/usr/tabsTAXI_TABSTRIP_OVERVIEW/tabpT\\01/ssubSUBSCREEN_BODY:SAPMV45A:4400/subSUBSCREEN_TC:SAPMV45A:4900/tblSAPMV45ATCTRL_U_ERF_AUFTRAG/txtKOMV-KBETR"
        
        try:
            for i in range(len(material)):
                mat_id = f"{mat_txt}[1,{i}]"
                qty_id = f"{qty_txt}[3,{i}]"
                amt_id = f"{amt_txt}[15,{i}]"
                
                print(f"Material ID: {mat_id}")
                print(f"Quantity ID: {qty_id}")
                print(f"Amount ID: {amt_id}")
                
                print(f"Material: {material[i]}")
                print(f"Quantity: {quantity[i]}")
                print(f"Amount: {amount[i]}")

                # Check if element exists before interacting
                if self.session.findById(mat_id) is not None:
                    self.session.findById(mat_id).text = material[i]
                else:
                    print(f"Material field not found: {mat_id}")
                    continue  # Skip to the next iteration

                if self.session.findById(qty_id) is not None:
                    self.session.findById(qty_id).text = quantity[i]
                else:
                    print(f"Quantity field not found: {qty_id}")
                    continue  # Skip to the next iteration

                if self.session.findById(amt_id) is not None:
                    self.session.findById(amt_id).text = amount[i]
                else:
                    print(f"Amount field not found: {amt_id}")
                    continue  # Skip to the next iteration

                # Process the entry
                self.session.findById("wnd[0]").sendVKey(0)
                # self.exceed_quantity_handling(error_id)
                # time.sleep(2)  # Ensure there is enough time between interactions
                # self.quantity_handling(window_id, text, button_id1, button_id2)
                # time.sleep(2)

        except Exception as e:
            print(f"An error occurred: {e}")

    def picked_qty_loc_select(self, picked_qty, location):
        picked_txt = "wnd[0]/usr/tabsTAXI_TABSTRIP_OVERVIEW/tabpT\\02/ssubSUBSCREEN_BODY:SAPMV50A:1104/tblSAPMV50ATC_LIPS_PICK/txtLIPSD-PIKMG"
        location_txt ="wnd[0]/usr/tabsTAXI_TABSTRIP_OVERVIEW/tabpT\\02/ssubSUBSCREEN_BODY:SAPMV50A:1104/tblSAPMV50ATC_LIPS_PICK/ctxtLIPS-LGORT"
        try:
            for i in range(len(picked_qty)):
                picked_id = f"{picked_txt}[7,{i}]"
                location_id =f"{location_txt}[3,{i}]"
                print(picked_id)
                print(location_id)
                print(picked_qty[i])
                print(location[i])
                self.session.findById(picked_id).text = picked_qty[i]
                self.session.findById(location_id).text = location[i]
                self.session.findById("wnd[0]").sendVKey(0)                         

        except Exception as e:
            print(e)

    def document_entry(self, doc_no):
        doc_txt = "wnd[0]/usr/tblSAPMV60ATCTRL_ERF_FAKT/ctxtKOMFK-VBELN"
        try:
            for i in range(len(doc_no)):
                picked_id = f"{doc_txt}[0,{i}]"
                print(picked_id)
                print(doc_no[i])
                self.session.findById(picked_id).text = doc_no[i]
                self.session.findById("wnd[0]").sendVKey(0)                         

        except Exception as e:
            print(e)


    def exceed_quantity_handling(self, error_id):
        try:
            status = self.session.findById(error_id).text
            print(status)
            status_split = status.split()
            status_splits = status_split[:-1]
            status_text = [status_split for status_split in status_splits if not status_split.isnumeric()]  # Fix the index here
            status1 = ' '.join(status_text)
            print(status1)
            if status1 == "Reorder point for item has been exceeded:":
                found = True
                self.session.findById("wnd[0]").sendVKey(0)
                time.sleep(5)
                return status
            else:
                print(status)  # Fix the variable name here
        except Exception as e:
            print(f"Error: {e}")

    def Extract_number(self, status_id):
        try:
            status = self.session.findById(status_id).text
            print(status)
            pattern = r'\b\d+\b'
            print(pattern)
            order_numbers = re.findall(pattern, status)
            print(order_numbers)
            if order_numbers:
                order_number = order_numbers[0]
                print("Order Number:", order_number)
                return order_number
            else:
                print("No order number found.")
        except Exception as e:
            print(f"Error: {e}")

    def close_window(self, window_id):
        self.session.findById(window_id).close()
        
    def to_upper(self, value):
        text = value.upper()
        print(text)
        return text
    def Delete_all_profile(self):
        try:
            self.session.findById("wnd[0]/usr/tabsTABSTRIP1/tabpPROF/ssubMAINAREA:SAPLSUID_MAINTENANCE:1103/cntlG_PROFILES_CONTAINER/shellcont/shell").setCurrentCell(-1, "")
            self.session.findById("wnd[0]/usr/tabsTABSTRIP1/tabpPROF/ssubMAINAREA:SAPLSUID_MAINTENANCE:1103/cntlG_PROFILES_CONTAINER/shellcont/shell").selectAll()
            self.session.findById("wnd[0]/usr/tabsTABSTRIP1/tabpPROF/ssubMAINAREA:SAPLSUID_MAINTENANCE:1103/cntlG_PROFILES_CONTAINER/shellcont/shell").pressToolbarButton("DEL_LINE")
        except:
            return  []
        
    def change_of_process(self, window_id, button_id):
        text = self.session.findById(window_id).text
        text_split = text.split()
        window = " ".join(text_split[:-1])
        if window == "Change processor of SAP Note":
            self.session.findById(button_id).press()

    def create_transport(self, window_id, create_button, text, finish_btn,):
        window = self.session.findById(window_id).Text
        if window == "Prompt for local Workbench request":
            transport = self.session.findById("wnd[1]/usr/ctxtKO008-TRKORR").Text
            if transport == "":
                self.session.findById(create_button).press()
                self.session.findById("wnd[2]/usr/txtKO013-AS4TEXT").Text = text
                self.session.findById(finish_btn).press()
                self.session.findById("wnd[1]/tbar[0]/btn[0]").press()
            else:
                self.session.findById("wnd[1]/tbar[0]/btn[0]").press()

    def get_license_product(self, element_id):
        license = self.session.findById(element_id).Text
        if license != "":
            license_split = license.split('_')
            product = license_split[0]
            return product
        else:
            return None
    def calculate_date_difference(self, given_date_str):
        given_date = datetime.strptime(given_date_str, '%d.%m.%Y').date()
        current_date = datetime.now().date()
        date_difference = (given_date - current_date).days
        return date_difference
    
    def write_value_to_excel(self, file_path, sheet_name, cell, value):
        workbook = load_workbook(file_path)
        sheet = workbook[sheet_name]
        try:
            sheet[cell] = float(value)
        except ValueError:
            sheet[cell] = str(value)
        workbook.save(file_path)
        workbook.close()
 
    def create_empty_excel(self, file_path):
        sheet_name = os.path.splitext(os.path.basename(file_path))[0]
        workbook = Workbook()
        workbook.active.title = sheet_name
        workbook.save(file_path)

    def excel_to_json(self, excel_file, json_file):
        # Read the Excel file
        df = pd.read_excel(excel_file, engine='openpyxl')
        # Convert Timestamp objects to strings
        for column in df.select_dtypes(['datetime']):
            df[column] = df[column].astype(str)
        # Convert the DataFrame to a dictionary
        data = df.to_dict(orient='records')
        # Write the dictionary to a JSON file
        with open(json_file, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=4)
        # Read the JSON file after writing it
        with open(json_file, 'r', encoding='utf-8') as f:
            json_data = json.load(f)
        return json_data
    
    def delete_specific_file(self, file_path):
            try:
                if os.path.exists(file_path):
                    os.remove(file_path)
                else:
                    print(f"The file '{file_path}' does not exist.")
            except Exception as e:
                print(f"An error occurred: {e}")
    
    def process_excel(self, file_path, sheet_name, column_index=None):
        df = pd.read_excel(file_path, sheet_name=sheet_name, header=None)
        if column_index is not None:
            try:
                column_index = int(column_index)  # Ensure column_index is an integer
            except ValueError:
                print("Invalid column index provided. Please provide a valid integer.")
                return
            if 0 <= column_index < df.shape[1]:
                df.drop(df.columns[column_index], axis=1, inplace=True)
            else:
                print(f"Column index {column_index} is out of bounds.")
                return
        df = df.applymap(lambda x: x.strip() if isinstance(x, str) else x)
        df.dropna(how='all', inplace=True)
        df.dropna(axis=1, how='all', inplace=True)
        if df.iloc[0].isnull().all(): 
            new_header = df.iloc[1]  
            df = df[2:]  # Remove first two rows
        else:
            new_header = df.iloc[0]  # Use first row as header
            df = df[1:]  # Remove first row
        df.columns = new_header
        df.reset_index(drop=True, inplace=True)
        try:
            # Write the modified DataFrame back to the Excel file
            with pd.ExcelWriter(file_path, engine='openpyxl', mode='w') as writer:
                df.to_excel(writer, index=False, sheet_name=sheet_name)
            print(f"Processed Excel sheet '{sheet_name}' has been updated in: {file_path}")
        except Exception as e:
            print(f"Error writing to Excel: {e}")

    def clean_list(self, data):
        cleaned_data = [item.strip() for item in data if item.strip()]
        return cleaned_data
    
    def read_table_column(self, table_id, column_index):
        """Reads the value of cell from selected column using column id
        """
        
        table = self.session.findById(table_id)
        row_count = self.session.findById(table_id).rowCount
       

        column_values = []
        for row in range(row_count):
            cell_value = self.get_cell_value(table_id,row, column_index)
            column_values.append(cell_value)

        return column_values
    
    def select_document_on_text(self, control_id, column_name, search_text):
        try:
            control = self.session.findById(control_id)
            row_count = control.RowCount  
            print(f"Total Rows: {row_count}")
            
            for row in range(row_count):
                print(f"Checking Row: {row}")
                cell_value = control.GetCellValue(row, column_name)  
                print(f"Cell Value in column '{column_name}': {cell_value}")
                
                if search_text in cell_value:
                    print(f"Text '{search_text}' found in row {row}")
                    
                    # Focus on the row before clicking (if required)
                    control.selectedRows = row  # Select the row (if applicable)
                    control.currentCellRow = row  # Set the current row to focus
                    
                    # Try different methods for clicking
                    control.doubleClickCell(row, column_name)  # Double click the cell
                    print(f"Double-clicked on row {row}")
                    
                    # Return the row where the text was found
                    return row  
                
                else:
                    print("Text not found in this row")
                    
            # If text is not found in any row
            return f"Text '{search_text}' not found"
        
        except Exception as e:
            return f"Error: {e}"
        
    def get_sap_table_value(self, table_id, row_num, column_id):
        # Get Sap Table Value    table_id=wnd[0]/usr/cntlGRID1/shellcont/shell    row_num=${row_num}    column_id=BELNR
        """
        Retrieves the value from a specific cell in the SAP table.

        :param table_id: The ID of the SAP table element
        :param row_num: The row number to get the value from
        :param column_id: The ID of the column to retrieve (e.g., "BELNR")
        :return: The value from the specified cell
        """
        try:
            table = self.session.findById(table_id)
            table.currentCellRow = row_num
            cell_value = table.getCellValue(row_num, column_id)
            return cell_value  
        except com_error as e:
            raise ValueError(f"Error retrieving value from SAP table: {e}")
        
    def select_row(self, table_id, row_number):
        # Select Row    table_id=wnd[0]/usr/cntlGRID1/shellcont/shell    row_number=${row_num}
        """
        Selects a specific row in the SAP table.

        :param table_id: The ID of the SAP table element
        :param row_number: The row number to select (0-based index)
        """
        try:
            table = self.session.findById(table_id)
            table.clearSelection()
            table.selectedRows = str(row_number)
            table.currentCellRow = row_number      
        except com_error as e:
            raise ValueError(f"Error selecting row {row_number} in SAP table: {e}")
        
    def expand_sap_shell_node(self, table_shell, row_number, column):
        try:
            row_number = int(row_number)
            if 0 <= row_number <= 9:
                row_identifier = f"{' ' * 10}{row_number}"  # 10 spaces for single digit
            elif 10 <= row_number <= 99:
                row_identifier = f"{' ' * 9}{row_number}"   # 9 spaces for double digits
            elif 100 <= row_number <= 999:
                row_identifier = f"{' ' * 8}{row_number}"   # 8 spaces for triple digits
            else:
                raise ValueError("Row number out of range. Must be between 0 and 999.")
            tree_id = table_shell 
            element = self.session.findById(tree_id)
            element.selectItem(row_identifier, column)
            element.expandNode(row_identifier)
        except Exception as e:
            raise Exception(f"Failed to expand node in SAP shell: {str(e)}")
        
    def get_sap_shell_item_value(self, table_shell, row_number, column):
        try:
            row_number = int(row_number)
            if 0 <= row_number <= 9:
                row_identifier = f"{' ' * 10}{row_number}"  # 10 spaces for single digit
            elif 10 <= row_number <= 99:
                row_identifier = f"{' ' * 9}{row_number}"   # 9 spaces for double digits
            elif 100 <= row_number <= 999:
                row_identifier = f"{' ' * 8}{row_number}"   # 8 spaces for triple digits
            else:
                raise ValueError("Row number out of range. Must be between 0 and 999.")
            element = self.session.findById(table_shell)
            return element.getItemText(row_identifier, column)
        except Exception as e:
            raise Exception(f"Failed to retrieve value from SAP shell: {str(e)}")
        
    def double_click_sap_shell_item(self, table_shell, row_number, column):
        try:
            row_number = int(row_number)
            if 0 <= row_number <= 9:
                row_identifier = f"{' ' * 10}{row_number}"  # 10 spaces for single digit
            elif 10 <= row_number <= 99:
                row_identifier = f"{' ' * 9}{row_number}"   # 9 spaces for double digits
            elif 100 <= row_number <= 999:
                row_identifier = f"{' ' * 8}{row_number}"   # 8 spaces for triple digits
            else:
                raise ValueError("Row number out of range. Must be between 0 and 999.")
            element = self.session.findById(table_shell)
            element.doubleClickItem(row_identifier, column)
        except Exception as e:
            raise Exception(f"Failed to double-click item in SAP shell: {str(e)}")
        
    def get_open_items_1(self, status_id):
        try:
            status = self.session.findById(status_id).Text
            # Regular expression to match all numbers in the text
            pattern = r"\d+"
            matches = re.findall(pattern, status)  # Find all matches
            if matches:
                # Convert matches to integers or return them as strings, depending on your requirement
                open_items = [int(item) for item in matches]
                return open_items
            else:
                print("No numeric values found")
                return []
        except Exception as e:
            return f"Error: {str(e)}"
        
    def Fbl1n_arrange(self):
        try:
            self.session.findById("wnd[1]/usr/tabsTS_LINES/tabpLI01/ssubSUB810:SAPLSKBH:0810/tblSAPLSKBHTC_WRITE_LIST").getAbsoluteRow(9).selected = True
            self.session.findById("wnd[1]/usr/tabsTS_LINES/tabpLI01/ssubSUB810:SAPLSKBH:0810/tblSAPLSKBHTC_WRITE_LIST").getAbsoluteRow(10).selected = True
            self.session.findById("wnd[1]/usr/tabsTS_LINES/tabpLI01/ssubSUB810:SAPLSKBH:0810/tblSAPLSKBHTC_WRITE_LIST/txtGT_WRITE_LIST-SELTEXT[0,9]").setFocus()
            self.session.findById("wnd[1]/usr/tabsTS_LINES/tabpLI01/ssubSUB810:SAPLSKBH:0810/tblSAPLSKBHTC_WRITE_LIST/txtGT_WRITE_LIST-SELTEXT[0,9]").caretPosition = 0
            self.session.findById("wnd[1]/usr/btnAPP_FL_SING").press()
            self.session.findById("wnd[1]/usr/tabsTS_LINES/tabpLI01/ssubSUB810:SAPLSKBH:0810/tblSAPLSKBHTC_WRITE_LIST").getAbsoluteRow(0).selected = True
            self.session.findById("wnd[1]/usr/btnAPP_FL_SING").press()
        except:
            return[]
        
    def process_excel(self, file_path, sheet_name, column_index=None, row_indices=None):
        try:
            workbook = load_workbook(filename=file_path)
            if sheet_name not in workbook.sheetnames:
                print(f"Sheet '{sheet_name}' not found in {file_path}.")
                return
            sheet = workbook[sheet_name]
            df = pd.DataFrame(sheet.values)
            if df.empty:
                print(f"The sheet '{sheet_name}' is empty.")
                return
            
            if column_index is not None:
                try:
                    column_index = [int(idx) for idx in column_index] if isinstance(column_index, list) else [int(column_index)]
                    for col in column_index:
                        if 0 <= col < df.shape[1]:
                            df.drop(columns=[col], inplace=True)
                        else:
                            print(f"Column index {col} is out of bounds.")
                            return
                except ValueError:
                    print("Invalid column index provided. Please provide valid integers.")
                    return
            
            df = df.apply(lambda x: x.str.strip() if x.dtype == "object" else x)
            df.dropna(how='all', inplace=True)
            df.dropna(axis=1, how='all', inplace=True)
            

            if df.iloc[0].isnull().all(): 
                new_header = df.iloc[1]  # Use second row as header
                df = df[2:]  # Remove first two rows
            else:
                new_header = df.iloc[0]  # Use first row as header
                df = df[1:]  # Remove first row
            df.columns = new_header
            df.reset_index(drop=True, inplace=True)

            if row_indices:
                try:
                    row_indices = [int(idx) for idx in row_indices] if isinstance(row_indices, list) else [int(row_indices)]
                    df.drop(index=row_indices, inplace=True)
                except Exception as e:
                    print(f"Error dropping rows: {e}")
                    return
            for row_idx, row in enumerate(df.itertuples(index=False, name=None), start=1):
                for col_idx, value in enumerate(row, start=1):
                    sheet.cell(row=row_idx + 1, column=col_idx, value=value)  # Starting from the first row of data

            workbook.save(file_path)
            workbook.close()
            print(f"Processed Excel sheet '{sheet_name}' has been updated in: {file_path}")
            
        except Exception as e:
            print(f"An error occurred: {e}")

    def click_node_link(self, tree_id, link_id1, link_id2, link_id3, link_id4, link_id5):
        """Selects a link of a TableTreeControl 'tree_id' which is contained within a shell object.
        
        Use the Scripting tracker recorder to find the 'link_id1' and 'link_id2' of the link to select.
        """
            
        self.session.findById(tree_id).expandNode(link_id1)
        self.session.findById(tree_id).topNode = link_id2
        self.session.findById(tree_id).expandNode(link_id3)
        self.session.findById(tree_id).selectItem(link_id4, link_id5)
        self.session.findById(tree_id).ensureVisibleHorizontalItem(link_id4, link_id5)
        self.session.findById(tree_id).topNode = link_id2
        self.session.findById(tree_id).clickLink(link_id4, link_id5)

    def sap_tcode_usmm_reg5(self):
        self.session.findbyId("wnd[0]/usr/cntlSLIM_USER_CONTAINER/shellcont/shell").pressToolbarContextButton("&MB_EXPORT")
        self.session.findbyId("wnd[0]/usr/cntlSLIM_USER_CONTAINER/shellcont/shell").selectContextMenuItem("&PC")

    def sap_tcode_usmm_reg5_filter(self):
        self.session.findbyId("wnd[0]/usr/cntlSLIM_USER_CONTAINER/shellcont/shell").selectcolumn("LICENSE_TYPE")
        self.session.findbyId("wnd[0]/usr/cntlSLIM_USER_CONTAINER/shellcont/shell").pressToolbarButton("&MB_FILTER")

    def mcr_report_pdf(self, images_directory):
        # Read data from the Excel file
        df = pd.read_excel("C:\\SAP_Testing\\SAPtesting\\Execution\\MCR_Input.xlsx")

        # Create a new Word document
        doc = docx.Document()

        section = doc.sections[-1]
        section.orientation = WD_ORIENT.LANDSCAPE

        # Add a Title to the document 
        doc.add_heading('Monthly Compliance Report', 0)

        # Add a table to the Word document
        table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1], style="Table Grid")
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Add column headers to the table
        for j in range(df.shape[1]):
            table.cell(0, j).text = df.columns[j]

        # Add data from the DataFrame to the table
        for i in range(df.shape[0]):
            for j in range(df.shape[1]):
                table.cell(i + 1, j).text = str(df.values[i, j])

        print("Data from test.xlsx has been successfully added to output.docx as a table.")

        widths = (Inches(0.5), Inches(1.0),  Inches(0.8), Inches(1.6), Inches(4))
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width=width

        for i in range(df.shape[0]):
            cell_value = str(df.values[i,4])
            no_of_images=list(cell_value.split(','))
            for j in range(len(no_of_images)):
                cell_1=table.cell(i+1,4)
                cell_1.add_paragraph().add_run().add_picture((os.path.join(images_directory, str(no_of_images[j]))), width=Inches(3.5), height=Inches(2.5))
        # Save the Word document
        doc.save("MCR_output.docx")
        #Convert to PDF
        convert("MCR_output.docx", "MCR_output_New.pdf")
        #doc.Close()
   
    def image_resize(self, screenshots_directory):
        # Define the directory path where your images are stored
        #directory = "C://RobotFramework//SAPtesting//Output//pabot_results//0"

        # Define the new size you want for your images
        new_size = (800, 600)
        keyword="Req5"
        # Loop through all the files in the directory
        for filename in os.listdir(screenshots_directory):
            if filename.endswith('.jpg') or filename.endswith('.png'):  # Add or remove file types as needed
                img_path = os.path.join(screenshots_directory, filename)
                if keyword in filename:
                    with Image.open(img_path) as image:
                    # Resize the image
                        width, height = image.size
                        # Setting the points for cropped image
                        left = 5
                        top = height / 12
                        right = 1800
                        bottom = 4 * height / 4
                        # Cropped image of above dimension
                        # (It will not change original image)
                        image_resize = image.crop((left, top, right, bottom))
                else:
                    with Image.open(img_path) as image:
                        # Resize the image
                        #img = img.resize(new_size, Image.ANTIALIAS)
                        width, height = image.size
                        # Setting the points for cropped image
                        left = 5
                        top = height / 12
                        right = 1200
                        bottom = 4 * height / 4

                        # Cropped image of above dimension
                        # (It will not change original image)
                        image_resize = image.crop((left, top, right, bottom))
                
                # Define the new filename
                new_filename = f"resized_{filename}"
                new_img_path = os.path.join(screenshots_directory, new_filename)
                
                # Save the resized image with the new name
                image_resize.save(new_img_path)

        print('All images have been resized.')

    def req7_usernames_extract(self, file_loc):
        df = pd.read_excel(file_loc, header=None, usecols=[0])

        with open("C:\\tmp\\output_req7.txt", 'w') as f:
            dfAsString = df.to_string(header=None, index=False)
            f.write(dfAsString)

    def copy_images(self, source_dir, target_dir):
        """
        Copy valid image files from source_dir to target_dir.
        Supported formats are '.jpeg', '.jpg', '.png', '.gif', '.bmp', '.tiff', '.webp'.
        """
        # Ensure the target directory exists
        if not os.path.exists(target_dir):
            os.makedirs(target_dir)

        # Supported image formats by PIL (Pillow)
        image_formats = ('.jpeg', '.jpg', '.png', '.gif', '.bmp', '.tiff', '.webp')

        # Iterate over files in the source directory
        for file_name in os.listdir(source_dir):
            file_path = os.path.join(source_dir, file_name)

            # Check if it's a file and if it has a valid image format
            if os.path.isfile(file_path):
                try:
                    with Image.open(file_path) as img:  # This will fail if the file is not a valid image
                        if file_name.lower().endswith(image_formats):
                            target_path = os.path.join(target_dir, file_name)
                            shutil.copy(file_path, target_path)
                            print(f"Copied: {file_name}")
                except Exception as e:
                    print(f"Skipped: {file_name} - Not a valid image file or format not supported. Error: {e}")

    def convert_xls_to_xlsx(self, xls_file, xlsx_file):
        """Convert .xls file to .xlsx format."""
        excel = win32.DispatchEx('Excel.Application')
        wb = None  # Initialize wb to None
        try:
            # Open the .xls file
            wb = excel.Workbooks.Open(xls_file)
            # Save it as .xlsx
            wb.SaveAs(xlsx_file, FileFormat=51)  # 51 corresponds to .xlsx
            print(f"Successfully converted {xls_file} to {xlsx_file}.")
        except Exception as e:
            print(f"Error: {e}")
        finally:
            # Close the workbook if it was successfully opened
            if wb:
                wb.Close()
            if excel:
                excel.Quit()

    def remove_rows_before_start_row(self, file_path, sheet_name, start_row):
        """
        Removes all rows before the specified start row in an Excel sheet.

        :param file_path: Path to the Excel file.
        :param sheet_name: Name of the sheet to modify.
        :param start_row: The row number to start from (all rows above this will be removed).
        """
        from openpyxl import load_workbook

        # Ensure start_row is an integer
        start_row = int(start_row)

        workbook = load_workbook(file_path)
        sheet = workbook[sheet_name]

        # Delete rows before start_row
        for _ in range(start_row - 1):  # Start row number is 1-based
            sheet.delete_rows(1)  # Always delete the first row

        workbook.save(file_path)
        workbook.close()
        return f"All rows before row {start_row} have been removed."
    
    def compare_and_include_query_data(self, security_file, query_file, output_file):
        """
        Compares rows from 'security_file' with 'query_file' and includes 'User Type' and 'Department'
        from both files in the output, along with marking missing rows or mismatched data.

        Args:
            security_file (str): Path to the SAP Security Users Excel file.
            query_file (str): Path to the SAP_QUERY Excel file.
            output_file (str): Path to save the comparison results.
        """
        # Load the data from both files
        security_df = pd.read_excel(security_file, sheet_name=0).fillna("Missing")
        query_df = pd.read_excel(query_file, sheet_name=0).fillna("Missing")

        # Standardize column names
        security_df.columns = security_df.columns.str.strip().str.lower()
        query_df.columns = query_df.columns.str.strip().str.lower()

        # Rename columns for consistency
        security_df.rename(columns={'user name': 'user_name', 'user type': 'user_type', 'department': 'department'}, inplace=True)
        query_df.rename(columns={'user name': 'user_name', 'user type': 'user_type', 'department': 'department'}, inplace=True)

        # Initialize results list
        results = []

        # Compare each row in security_df with query_df
        for _, row in security_df.iterrows():
            user_name = row['user_name']
            user_type_security = row['user_type']
            department_security = row['department']

            # Check if the username exists in query_df
            matched_row = query_df[query_df['user_name'] == user_name]

            if matched_row.empty:
                # User is completely missing
                results.append([
                    user_name, user_type_security, department_security,
                    "Missing", "Missing", "Missing User"
                ])
            else:
                # Extract query data
                user_type_query = matched_row.iloc[0]['user_type']
                department_query = matched_row.iloc[0]['department']

                # Check for mismatches
                status_user_type = "Matched" if user_type_query == user_type_security else "Mismatch"
                status_department = "Matched" if department_query == department_security else "Mismatch"

                results.append([
                    user_name, user_type_security, department_security,
                    user_type_query, department_query,
                    "User exists but has mismatched data" if "Mismatch" in (status_user_type, status_department) else "Fully Matched"
                ])

        # Convert results to DataFrame
        results_df = pd.DataFrame(results, columns=[
            'User Name', 'User Type (Security)', 'Department (Security)',
            'User Type (Query)', 'Department (Query)', 'Notes'
        ])

        # Save results to an Excel file
        with pd.ExcelWriter(output_file, engine='openpyxl') as writer:
            results_df.to_excel(writer, index=False, sheet_name='SAP_QUERY')

        # Apply formatting with openpyxl
        wb = Workbook()
        ws = wb.active
        ws.title = "Comparison Results"

        # Write headers with styling
        headers = [
            'User Name', 'User Type (Security)', 'Department (Security)',
            'User Type (Query)', 'Department (Query)', 'Notes'
        ]
        for col_num, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_num, value=header)
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="ADD8E6", end_color="ADD8E6", fill_type="solid")

        # Write rows
        for row_idx, row in enumerate(results, start=2):
            for col_idx, value in enumerate(row, start=1):
                cell = ws.cell(row=row_idx, column=col_idx, value=value)

                # Highlight missing or mismatched data
                if value == "Missing User" or "Mismatch" in str(value):
                    cell.fill = PatternFill(start_color="FFCCCC", end_color="FFCCCC", fill_type="solid")
                    cell.font = Font(bold=True)

        # Save formatted workbook
        wb.save(output_file)
        print(f"Comparison completed! Results saved to: {output_file}")

    def click_node_link_1(self, tree_id, link_id6, link_id7, link_id8, link_id9, link_id10):
            """Selects a link of a TableTreeControl 'tree_id' which is contained within a shell object.
        
            Use the Scripting tracker recorder to find the 'link_id1' and 'link_id2' of the link to select.
            """
            
            self.session.findById(tree_id).expandNode(link_id6)
            self.session.findById(tree_id).topNode = link_id7
            self.session.findById(tree_id).expandNode(link_id8)
            self.session.findById(tree_id).selectItem(link_id9, link_id10)
            self.session.findById(tree_id).ensureVisibleHorizontalItem(link_id9, link_id10)
            self.session.findById(tree_id).topNode = link_id7
            self.session.findById(tree_id).clickLink(link_id9, link_id10)

    def click_node_link_3(self, tree_id, link_id1, link_id2, link_id13, link_id14):
            """Selects a link of a TableTreeControl 'tree_id' which is contained within a shell object.
        
            Use the Scripting tracker recorder to find the 'link_id1' and 'link_id2' of the link to select.
            """
    
            self.session.findById(tree_id).expandNode(link_id1)
            self.session.findById(tree_id).topNode = link_id2
            self.session.findById(tree_id).selectItem(link_id13, link_id14)
            self.session.findById(tree_id).ensureVisibleHorizontalItem(link_id13, link_id14)
            self.session.findById(tree_id).clickLink(link_id13, link_id14)

    def extract_columns(self, file1, sheet1, col1_index, skiprows, file2, sheet2, col2_index, output_file, header1, header2):
        # Ensure the column indices are integers in case they are passed as strings
        col1_index = int(col1_index)
        col2_index = int(col2_index)
        skiprows = int(skiprows)  # Convert skiprows to integer if it's a string

        # Read data from the first file, skipping the specified number of rows
        df1 = pd.read_excel(file1, sheet_name=sheet1, skiprows=skiprows)  # Skip specified rows
        df2 = pd.read_excel(file2, sheet_name=sheet2)  # Read data from the second Excel file

        # Validate that the column indices are within bounds
        if col1_index >= len(df1.columns) or col2_index >= len(df2.columns):
            raise ValueError("Column index is out of bounds for one of the files.")

        # Create a new DataFrame for the output
        result_df = pd.DataFrame()

        # Extract data using column indices, dropping NaN values to handle empty rows
        result_df[header1] = df1.iloc[:, col1_index].dropna().reset_index(drop=True)
        result_df[header2] = df2.iloc[:, col2_index].dropna().reset_index(drop=True)

        # Write the result to a new Excel file
        result_df.to_excel(output_file, index=False)

        return "Columns extracted successfully."
    
    def compare_columns(self, input_file, col1_name, col2_name, comparison_col_name):
    # Read the combined Excel file
        df = pd.read_excel(input_file)

        # Initialize a list to hold the comparison results
        comparison_results = []

        # Convert the second column to a set for faster lookup
        col2_set = set(df[col2_name])  # Use the column name passed as an argument

        # Iterate through the first column and check if each username exists in the second column
        for col1_value in df[col1_name]:  # Use the column name passed as an argument
            if col1_value in col2_set:
                comparison_results.append(col1_value)  # Add the username if it exists in Column2
            else:
                comparison_results.append("")  # Keep it empty if it doesn't exist

        # Add the comparison results to the DataFrame in the specified comparison column
        df[comparison_col_name] = comparison_results  # Use the comparison column name passed as an argument

        # Write the updated DataFrame back to the same Excel file
        df.to_excel(input_file, index=False)

        return "Comparison completed successfully."
    
    def Matched_columns(self, input_file, col1_name, col2_name):
        # Configure logging to write to both file and console
        logging.basicConfig(level=logging.INFO, 
                            format='%(asctime)s - %(levelname)s - %(message)s', 
                            handlers=[
                                logging.FileHandler("matching.log"),
                                logging.StreamHandler()  # This logs to the console
                            ])
        
        # Read the combined Excel file
        df = pd.read_excel(input_file)

        # Initialize lists to hold the comparison results
        matched_results = []
        not_matched_results = []

        # Convert the second column to a set for faster lookup
        col2_set = set(df[col2_name])  # Use the column name passed as an argument

        # Iterate through the first column and check if each username exists in the second column
        for col1_value in df[col1_name]:  # Use the column name passed as an argument
            if col1_value in col2_set:
                matched_results.append(col1_value)  # Add the username if it exists in Column2
                not_matched_results.append("")  # Keep it empty for matched users
            else:
                matched_results.append("")  # Keep it empty for matched users
                not_matched_results.append(col1_value)  # Add the username if it does not exist in Column2

        # Add the results to the DataFrame in the respective columns
        df['Matched'] = matched_results
        df['Not Matched'] = not_matched_results

        # Write the updated DataFrame back to the same Excel file
        df.to_excel(input_file, index=False)

        # Check for "Not Matched" entries
        unauthorized_users = [user for user in not_matched_results if user]  # Filter out empty strings

        # Log and display messages based on the results
        if unauthorized_users:
            message = "The listed users are unauthorized."
        else:
            message = "All users are authorized."

        # Log the message to both console and log file
        logging.info(message)

        return message
    
    def compare_and_add_query_data(self, security_file, query_file, existing_file, new_sheet_name):
    
        # Load the data from both files
        security_df = pd.read_excel(security_file, sheet_name=0).fillna("Missing")
        query_df = pd.read_excel(query_file, sheet_name=0).fillna("Missing")

        # Standardize column names
        security_df.columns = security_df.columns.str.strip().str.lower()
        query_df.columns = query_df.columns.str.strip().str.lower()

        # Rename columns for consistency
        security_df.rename(columns={'user name': 'user_name', 'user type': 'user_type', 'department': 'department'}, inplace=True)
        query_df.rename(columns={'user name': 'user_name', 'user type': 'user_type', 'department': 'department'}, inplace=True)

        # Initialize results list
        results = []

        # Compare each row in security_df with query_df
        for _, row in security_df.iterrows():
            user_name = row['user_name']
            user_type_security = row['user_type']
            department_security = row['department']

            # Check if the username exists in query_df
            matched_row = query_df[query_df['user_name'] == user_name]

            if matched_row.empty:
                # User is completely missing
                results.append([
                    user_name, user_type_security, department_security,
                    "Missing", "Missing", "Missing User"
                ])
            else:
                # Extract query data
                user_type_query = matched_row.iloc[0]['user_type']
                department_query = matched_row.iloc[0]['department']

                # Check for mismatches
                status_user_type = "Matched" if user_type_query == user_type_security else "Mismatch"
                status_department = "Matched" if department_query == department_security else "Mismatch"

                results.append([
                    user_name, user_type_security, department_security,
                    user_type_query, department_query,
                    "User exists but has mismatched data" if "Mismatch" in (status_user_type, status_department) else "Fully Matched"
                ])

        # Convert results to DataFrame
        results_df = pd.DataFrame(results, columns=[
            'User Name', 'User Type (Security)', 'Department (Security)',
            'User Type (Query)', 'Department (Query)', 'Notes'
        ])

        # Append results as a new sheet to the existing Excel file
        try:
            # Load the workbook if it exists
            workbook = load_workbook(existing_file)
        except FileNotFoundError:
            # Create a new workbook if the file does not exist
            workbook = Workbook()
            if "Sheet" in workbook.sheetnames and len(workbook["Sheet"]["A"]) == 0:
                workbook.remove(workbook["Sheet"])

        # Check if the sheet already exists
        if new_sheet_name in workbook.sheetnames:
            raise ValueError(f"Sheet '{new_sheet_name}' already exists in the workbook.")

        # Add a new sheet to the workbook
        sheet = workbook.create_sheet(new_sheet_name)

        # Write data to the new sheet
        for row in dataframe_to_rows(results_df, index=False, header=True):
            sheet.append(row)

        # Apply formatting to the new sheet
        for col_num, cell in enumerate(sheet[1], start=1):
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="ADD8E6", end_color="ADD8E6", fill_type="solid")

        # Highlight mismatches or missing rows
        for row in sheet.iter_rows(min_row=2, max_row=sheet.max_row, min_col=1, max_col=sheet.max_column):
            for cell in row:
                if "Mismatch" in str(cell.value) or "Missing" in str(cell.value):
                    cell.fill = PatternFill(start_color="FFCCCC", end_color="FFCCCC", fill_type="solid")
                    cell.font = Font(bold=True)

        # Save the workbook
        workbook.save(existing_file)
        print(f"Comparison completed! Results saved to sheet '{new_sheet_name}' in '{existing_file}'.")

    def click_current_cell(self,element_id,cell_value):
            try:
                element =self.session.findById(element_id)
                element.currentCellColumn = cell_value
                element.clickCurrentCell()
            except Exception as e:
                print(f"Error: {e}")

    def Generate_Word(self, excel_directory, images_directory, doc_name):
    # Read data from the Excel file
        df = pd.read_excel(excel_directory)

        # Create a new Word document
        doc = Document()

        # Set the document orientation to landscape
        section = doc.sections[-1]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11.69)
        section.page_height = Inches(13.27)  # A4 landscape size
        
        # Adjust page margins
        section.top_margin = Inches(1.5)  # Increase the top margin to make room for header content
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        # Add a custom header
        header = section.header
        header_table = header.add_table(rows=1, cols=3, width = Inches(6))  # Create the table
        header_table.style = "Table Grid"  # Apply the style here
        header_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Adjust column widths
        header_table.columns[0].width = Inches(1.2)
        header_table.columns[1].width = Inches(1.6)
        header_table.columns[2].width = Inches(1.2)
        
        for cell in header_table.rows[0].cells:
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Column 1: Logo
        cell1 = header_table.cell(0, 0)
        paragraph = cell1.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_path = "C:\\tmp\\BCS.png"
        if os.path.exists(logo_path):
            paragraph.add_run().add_picture(logo_path, width=Inches(0.4), height=Inches(0.4))
        else:
            run = paragraph.add_run("Logo Missing")
            run.bold = True
        
        # Column 2: Title
        cell2 = header_table.cell(0, 1)
        title_paragraph = cell2.paragraphs[0]
        title_run = title_paragraph.add_run("Monthly Compliance Report")
        title_run.bold = True
        title_run.font.size = Pt(11)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Column 3: Version Details
        cell3 = header_table.cell(0, 2)
        version_paragraph = cell3.paragraphs[0]
        version_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        version_paragraph.add_run("Release Version: 0.1\n").bold = True
        version_paragraph.add_run("_" * 50 + "\n").font.size = Pt(8)
        version_paragraph.add_run("Guideline No: 0.1").bold = True
        
        # Add a footer
        footer = section.footer
        footer_paragraph = footer.add_paragraph("MIC")
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_paragraph.runs[0].italic = True
        
        # Add spacing after the header
        # spacer = doc.add_paragraph()
        # spacer_format = spacer.paragraph_format
        # spacer_format.space_after = Pt(12)

        # Add a data table
        data_table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
        data_table.style = "Table Grid"

        # Adjust column widths
        for idx, width in enumerate([Inches(1.0), Inches(1.5), Inches(1.2), Inches(2.0), Inches(3.5)]):
            for row in data_table.rows:
                row.cells[idx].width = width

        # Style the header row
        hdr_cells = data_table.rows[0].cells
        header_bg_color = "0066cc"
        for j in range(df.shape[1]):
            hdr_cells[j].text = df.columns[j]
            hdr_cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            hdr_cells[j]._element.get_or_add_tcPr().append(
                parse_xml(f'<w:shd {nsdecls("w")} w:fill="{header_bg_color}"/>')
            )

        # Add data to the table
        for i in range(df.shape[0]):
            for j in range(df.shape[1]):
                data_table.cell(i + 1, j).text = str(df.iloc[i, j])

        # Embed images into the last column
        for i in range(df.shape[0]):
            cell_value = str(df.iloc[i, -1])  # Assuming images are in the last column
            images = cell_value.split(',')
            for img_name in images:
                img_path = os.path.join(images_directory, img_name.strip())
                if os.path.exists(img_path):
                    cell = data_table.cell(i + 1, df.shape[1] - 1)
                    cell.add_paragraph().add_run().add_picture(img_path, width=Inches(3.5), height=Inches(2.5))

        # Save the Word document and convert to PDF
        doc.save(f"{doc_name}.docx")
        convert(f"{doc_name}.docx", f"{doc_name}.pdf")

    def files_clean_username(self, file_path, cleaned_file_path):
        
        df = pd.read_excel(file_path, header=None)

        # Find the row index where 'User Name' appears
        header_row_idx = df[df.eq("User Name").any(axis=1)].index[0]

        # Re-load the data with the correct header row
        df_cleaned = pd.read_excel(file_path, skiprows=header_row_idx)

        # Save the cleaned file
        df_cleaned.to_excel(cleaned_file_path, index=False)

    def files_clean_Bname(self, file_path, cleaned_file_path):
        
        df = pd.read_excel(file_path, header=None)

        # Find the row index where 'User Name' appears
        header_row_idx = df[df.eq("BNAME").any(axis=1)].index[0]

        # Re-load the data with the correct header row
        df_cleaned = pd.read_excel(file_path, skiprows=header_row_idx)

        # Rename 'User' column to 'User Name' (if it exists)
        df_cleaned.rename(columns={"BNAME": "User Name"}, inplace=True)

        # Save the cleaned file
        df_cleaned.to_excel(cleaned_file_path, index=False)

    def files_clean_Uname(self, file_path, cleaned_file_path):
        
        df_cleaned = pd.read_excel(file_path, header=None)

        # Find the row index where 'User Name' appears
        # header_row_idx = df[df.eq("UNAME").any(axis=1)].index[0]

        # Re-load the data with the correct header row
        # df_cleaned = pd.read_excel(file_path, skiprows=header_row_idx)

        # Rename 'User' column to 'User Name' (if it exists)
        df_cleaned.rename(columns={"UNAME": "User Name"}, inplace=True)

        # Save the cleaned file
        df_cleaned.to_excel(cleaned_file_path, index=False)

    def files_clean_Username_cleaned(self, file_path, cleaned_file_path):
        
        df_cleaned = pd.read_excel(file_path, header=None)

        # Find the row index where 'User Name' appears
        # header_row_idx = df[df.eq("UNAME").any(axis=1)].index[0]

        # Re-load the data with the correct header row
        # df_cleaned = pd.read_excel(file_path, skiprows=header_row_idx)

        # Rename 'User' column to 'User Name' (if it exists)
        # df_cleaned.rename(columns={"UNAME": "User Name"}, inplace=True)

        # Save the cleaned file
        df_cleaned.to_excel(cleaned_file_path, index=False)
        
    def files_clean_user(self, file_path, cleaned_file_path):
        
        df = pd.read_excel(file_path, header=None)

        # Find the row index where 'User' appears
        header_row_idx = df[df.eq("User").any(axis=1)].index[0]

        # Re-load the data with the correct header row
        df_cleaned = pd.read_excel(file_path, skiprows=header_row_idx)

        # Rename 'User' column to 'User Name' (if it exists)
        df_cleaned.rename(columns={"User": "User Name"}, inplace=True)

        # Save the cleaned file
        df_cleaned.to_excel(cleaned_file_path, index=False)
        
    def generate_extra_users_list(self, se16_filepath_cleaned, table_auth_filepath_cleaned, control_SAP_filepath_cleaned, auth_profiles_filepath_cleaned, exempted_users_file, output_file):
        
        # File paths
        #exempted_users_file = "C:\\MCR_Report_Files\\MCR_Exempted_Users1.xlsx"

        files_to_compare = [
            se16_filepath_cleaned,
            table_auth_filepath_cleaned,
            control_SAP_filepath_cleaned,
            auth_profiles_filepath_cleaned
        ]

        #output_file = "C:\\MCR_Report_Files\\Extra_Users_List_full.xlsx"

        # Load exempted users file with all sheets
        exempted_users_sheets = pd.read_excel(exempted_users_file, sheet_name=None)
        print

        # Dictionary to store extra users per file
        extra_users_dict = {}

        # Compare each file to its corresponding sheet
        for file_path, (sheet_name, exempted_users_df) in zip(files_to_compare, exempted_users_sheets.items()):
            try:
                # Load the current file
                output_df = pd.read_excel(file_path)

                # Ensure 'User Name' column exists in both files
                if ('User Name' or 'User') in exempted_users_df.columns and ('User Name' or 'User') in output_df.columns:
                    exempted_users = set(exempted_users_df['User Name' or 'User'])
                    output_users = set(output_df['User Name' or 'User'])

                    # Find extra users
                    extra_users_df = output_df[~output_df['User Name' or 'User'].isin(exempted_users)]

                    if not extra_users_df.empty:
                        # Extract only the file name without the path
                        file_name = os.path.basename(file_path).replace("Cleaned.xlsx", "")
                        sheet_name = f"Extra_Users_{file_name}"[:31]
                        
                        # Store results with safe sheet name
                        extra_users_dict[sheet_name] = extra_users_df
                    else:
                        print(f"✅ No extra users found in: {file_path} vs {sheet_name}")

            except Exception as e:
                print(f"⚠ Error processing {file_path} vs {sheet_name}: {e}")

        # Save all extra users in one Excel file with multiple sheets
        if extra_users_dict:
            with pd.ExcelWriter(output_file) as writer:
                for sheet, df in extra_users_dict.items():
                    df.to_excel(writer, sheet_name=sheet, index=False)

            print(f"✅ Comparison complete. Extra users saved in '{output_file}'.")
        else:
            print("✅ No extra users found in any comparison.")

    def compare(self, first_excel_path, second_excel_path, output_excel_path, 
                         first_sheet_name, second_sheet_name):
        """
        Compare two Excel sheets based on User Name, User Type, and Department, 
        and write mismatched rows to a new Excel file, highlighting missing Department values 
        and marking users with the user type 'S_TABU_DISP'.
        """
        # Load the Excel sheets
        first_df = pd.read_excel(first_excel_path, sheet_name=first_sheet_name)
        second_df = pd.read_excel(second_excel_path, sheet_name=second_sheet_name)

        # Ensure consistent column names
        first_df.columns = first_df.columns.str.strip().str.lower()
        second_df.columns = second_df.columns.str.strip().str.lower()

        # Rename columns to standardize
        first_df.rename(columns={
            'user name': 'user_name',
            'user type': 'user_type',
            'department': 'department'
        }, inplace=True)

        second_df.rename(columns={
            'user name': 'user_name',
            'user type': 'user_type',
            'department': 'department'
        }, inplace=True)

        # Merge data on 'user_name'
        merged_df = pd.merge(first_df, second_df, on='user_name', suffixes=('_first', '_second'))

        # Update column names in the merged DataFrame
        merged_df.rename(columns={
            'user_type_first': f'{first_sheet_name}_user_type',
            'user_type_second': f'{second_sheet_name}_user_type',
            'department_first': f'{first_sheet_name}_department',
            'department_second': f'{second_sheet_name}_department'
        }, inplace=True)

        # Find mismatched rows based on 'user_type' and 'department'
        mismatched_rows = merged_df[
            (merged_df[f'{first_sheet_name}_user_type'] != merged_df[f'{second_sheet_name}_user_type']) |
            (merged_df[f'{first_sheet_name}_department'] != merged_df[f'{second_sheet_name}_department'])
        ]

        # Select the relevant columns for output
        output_columns = [
            'user_name', 
            f'{first_sheet_name}_user_type', f'{second_sheet_name}_user_type', 
            f'{first_sheet_name}_department', f'{second_sheet_name}_department'
        ]

        # Filter the mismatched rows to only include the output columns
        mismatched_rows = mismatched_rows[output_columns]

        # Add missing markers for usertype or department
        mismatched_rows[f'{first_sheet_name}_user_type'] = mismatched_rows[f'{first_sheet_name}_user_type'].fillna('Missing: User Type')
        mismatched_rows[f'{second_sheet_name}_user_type'] = mismatched_rows[f'{second_sheet_name}_user_type'].fillna('Missing: User Type')
        mismatched_rows[f'{first_sheet_name}_department'] = mismatched_rows[f'{first_sheet_name}_department'].fillna('Missing: Department')
        mismatched_rows[f'{second_sheet_name}_department'] = mismatched_rows[f'{second_sheet_name}_department'].fillna('Missing: Department')

        # Check for user type 'S_TABU_DISP' in SAP Security Users sheet
        second_df['S_TABU_DISP_User'] = second_df['user_type'].apply(
            lambda x: 'Yes' if x == 'S_TABU_DISP' else 'No'
        )

        # Merge this information into mismatched rows for highlighting
        mismatched_rows = pd.merge(mismatched_rows, second_df[['user_name', 'S_TABU_DISP_User']], on='user_name', how='left')

        # Write mismatched rows to a new Excel file
        mismatched_rows.to_excel(output_excel_path, index=False)

        # Now apply bold formatting and highlight missing departments or S_TABU_DISP users using openpyxl
        wb = load_workbook(output_excel_path)
        ws = wb.active

        # Highlight "Missing: Department" cells in yellow
        missing_department_fill = PatternFill(start_color="FFFF99", end_color="FFFF99", fill_type="solid")
        s_tab_disp_fill = PatternFill(start_color="FFCCCC", end_color="FFCCCC", fill_type="solid")

        for row in ws.iter_rows(min_row=2, min_col=1, max_col=7, max_row=ws.max_row):  # Updated for 7 columns (added S_TABU_DISP_User)
            for cell in row:
                # Highlight missing department cells
                if isinstance(cell.value, str) and 'Missing: Department' in cell.value:
                    cell.fill = missing_department_fill
                    cell.font = Font(bold=True)  # Optional: Apply bold formatting

                # Highlight S_TABU_DISP user rows
                if isinstance(cell.value, str) and 'Yes' in cell.value:
                    cell.fill = s_tab_disp_fill
                    cell.font = Font(bold=True)  # Optional: Apply bold formatting

        wb.save(output_excel_path)

        print(f"Mismatched data has been written to {output_excel_path}")

    def generate_report_html(self, input_file, output_file, report_name):
        if not os.path.exists(input_file):
            print(f"Error: The input file '{input_file}' does not exist.")
            return

        pass_count = 0
        warn_count = 0
        current_time = datetime.datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S UTC')

        try:
            with open(input_file, 'r') as file:
                lines = file.readlines()
                for line in lines:
                    line = line.strip()
                    # Only count non-empty lines
                    if line:
                        if 'PASS' in line:
                            pass_count += 1
                        elif 'WARN' in line:
                            warn_count += 1

            html_content = f"""
            <!DOCTYPE html>
            <html lang="en">
            <head>
                <meta charset="UTF-8">
                <meta name="viewport" content="width=device-width, initial-scale=1.0">
                <title>MIC Report</title>
                <style>
                    body {{ font-family: Arial, sans-serif; }}
                    h1, h2 {{ text-align: center; }}
                    .pass {{ color: green; }}
                    .warn {{ color: red; }}
                    .pass-circle {{ 
                        width: 30px; height: 30px; background-color: green; 
                        border-radius: 50%; display: inline-block; margin-left: 10px;
                    }}
                    .warn-circle {{ 
                        width: 30px; height: 30px; background-color: red; 
                        border-radius: 50%; display: inline-block; margin-left: 10px;
                    }}
                    .header {{
                        background-color: rgb(21, 2, 166);
                        color: white;
                        padding: 20px 0;
                        display: flex;
                        flex-direction: column;
                        justify-content: center;
                        align-items: center;
                        position: relative;
                    }}
                    .header img {{
                        width: 165px;
                        margin-bottom: 10px;
                    }}
                    .content {{
                        display: flex;
                        justify-content: center;
                        align-items: center;
                        flex-direction: column;
                        margin-top: 20px;
                    }}
                    .title-text {{
                        text-align: center;
                        font-size: 12px;
                        margin-top: 20px;
                    }}
                    table {{
                        width: 80%;
                        border-collapse: collapse;
                        margin: 20px 0;
                    }}
                    th, td {{
                        padding: 8px;
                        border: 1px solid #ddd;
                        text-align: left;
                    }}
                    th {{ background-color: #f4f4f4; }}
                    .count-box {{
                        display: flex;
                        justify-content: center;
                        align-items: center;
                        gap: 20px;
                        margin: 20px 0;
                    }}
                    .count-item {{
                        display: flex;
                        align-items: center;
                        gap: 10px;
                        font-size: 18px;
                    }}
                </style>
            </head>
            <body>
                <div class="header">
                    <img src="https://symphony4cloud.com/static/media/logo-1.a84c61ed3670ed0777c8.png" alt="Logo">
                    <div class="title">
                        AI Powered Automation Report
                    </div>
                </div>

                <div class="title-text">
                    <h1> {report_name} </h1>
                    <h2>Report Generated: {current_time}</h2>
                </div>

                <div class="content">
                    <div class="count-box">
                        <div class="count-item"><strong>PASS Count:</strong> {pass_count} <span class="pass-circle"></span></div>
                        <div class="count-item"><strong>WARN Count:</strong> {warn_count} <span class="warn-circle"></span></div>
                    </div>

                    <table>
                        <tr>
                            <th>Seq</th>
                            <th>Message</th>
                        </tr>
            """
            counter = 1
            for line in lines:
                line = line.strip()
                if line:  # Ensure non-empty lines are processed
                    row_class = ""
                    if 'PASS' in line:
                        row_class = "pass"
                    elif 'WARN' in line:
                        row_class = "warn"
                    
                    # Add each line with the correct class based on 'PASS' or 'WARN'
                    html_content += f'<tr><td>{counter}</td><td class="{row_class}">{line}</td></tr>\n'
                    counter += 1

            html_content += f"""
                    </table>
                </div>

                <div class="footer">
                    <p>Report of &copy; {datetime.datetime.now().year}</p>
                </div>
            </body>
            </html>
            """
            output_dir = os.path.dirname(output_file)
            if not os.path.exists(output_dir):
                os.makedirs(output_dir)

            with open(output_file, 'w') as file:
                file.write(html_content)
            print(f"Report generated successfully: {output_file}")
        except Exception as e:
            print(f"Error processing the report: {e}")


    def generate_extra_users_list_all_nodes(self, file_list_path, exempted_users_file, output_file):
       
        import pandas as pd
 
        # File paths
        #exempted_users_file = "C:\\MCR_Report_Files\\MCR_Exempted_Users1.xlsx"
        #file_list_path = "C:\\MCR_Report_Files\\file_list.txt"
        #output_file = "C:\\MCR_Report_Files\\Extra_Users_List.xlsx"
 
        # Read the list of files from the text file
        with open(file_list_path, "r") as f:
            file_paths = [line.strip() for line in f if line.strip()]
 
        # Load the first sheet of the exempted users file
        exempted_users_sheets = pd.read_excel(exempted_users_file, sheet_name=None)
        exempted_users_list = list(exempted_users_sheets.values())
 
        # Dictionary to store extra users per file
        extra_users_dict = {}
 
        # Process each file in the list and compare it with the corresponding exempted sheet
        for idx, file_path in enumerate(file_paths):
            try:
                output_df = pd.read_excel(file_path)
                exempted_users_df = exempted_users_list[idx]  # Match file with corresponding exempted sheet
 
                # Ensure 'User Name' exists in both dataframes
                if 'User Name' in exempted_users_df.columns and 'User Name' in output_df.columns:
                    exempted_users = set(exempted_users_df['User Name'])
                    output_users = set(output_df['User Name'])
 
                    # Find extra users
                    extra_users_df = output_df[~output_df['User Name'].isin(exempted_users)]
 
                    if not extra_users_df.empty:
                        # sheet_name = f"Comparison_{idx+1}"
                        file_name = os.path.basename(file_path).replace("Cleaned.xlsx", "")
                        sheet_name = f"{file_name}"[:31]
                        extra_users_dict[sheet_name] = extra_users_df
                else:
                    print(f"⚠ 'User Name' column not found in {file_path}. Skipping...")
 
            except Exception as e:
                print(f"⚠ Error processing {file_path}: {e}")
 
        # Save results to an Excel file with multiple sheets
        if extra_users_dict:
            with pd.ExcelWriter(output_file) as writer:
                for sheet, df in extra_users_dict.items():
                    df.to_excel(writer, sheet_name=sheet[:31], index=False)
 
            print(f"✅ Comparison complete. Extra users saved in '{output_file}'.")
        else:
            print("⚠ No extra users found. Output file not created.")
    def extract_numeric(self, data):
        match = re.search(r'\d+', data)
        if match:
            return match.group()
        else:
            return data
    def rpa_dropdown_bin_conformation(self,tableshell,Button,MenuItem):
        try:
            self.session.findById(tableshell).pressToolbarContextButton(Button)
            self.session.findById(tableshell).selectContextMenuItem(MenuItem)
        except:
            return  []
    def rpa_selected_rows(self, tree_id, first_visible_row):
        try:
            self.session.findById(tree_id).currentCellRow = first_visible_row
            self.session.findById(tree_id).selectedRows = first_visible_row
        except Exception as e:
            print(f"Error: {e}")
    
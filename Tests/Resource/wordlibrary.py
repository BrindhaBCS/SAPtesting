from docx import Document
from robot.api.deco import keyword

class wordlibrary:

    @keyword
    def read_word_file(self, file_path):
        """Reads the content of a Word file and returns it as a string."""
        doc = Document(file_path)
        content = []
        for para in doc.paragraphs:
            content.append(para.text)
        return "\n".join(content)

    @keyword
    def write_word_file(self, file_path, content):
        """Writes the given content to a Word file."""
        doc = Document()
        doc.add_paragraph(content)
        doc.save(file_path)
